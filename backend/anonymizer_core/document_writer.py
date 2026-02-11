"""Document writer module for saving anonymized documents."""

import io
import re
import zlib
from pathlib import Path
from typing import Any
from copy import deepcopy
from datetime import datetime, date

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
import openpyxl
from openpyxl.styles import PatternFill

try:
    import pikepdf
    HAS_PIKEPDF = True
except ImportError:
    HAS_PIKEPDF = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

from .anonymizer import AnonymizationResult, Replacement


class DocumentWriter:
    """Writer for creating anonymized documents."""
    
    def __init__(self):
        # Placeholder image for removed logos
        self.logo_placeholder_text = "[Логотип удален]"
    
    def write_docx(
        self, 
        original_path: Path, 
        output_path: Path,
        result: AnonymizationResult,
        remove_images: bool = False
    ) -> Path:
        """
        Create an anonymized Word document.
        
        Args:
            original_path: Path to original document
            output_path: Path to save anonymized document
            result: AnonymizationResult with replacements
            remove_images: Whether to remove images (logos)
            
        Returns:
            Path to the saved document
        """
        doc = Document(original_path)
        
        # Build replacement map
        replacement_map = {r.original: r.anonymized for r in result.replacements}
        
        # Process paragraphs
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, replacement_map)
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, replacement_map)
        
        # Process headers and footers
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for paragraph in header.paragraphs:
                        self._replace_in_paragraph(paragraph, replacement_map)
            
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for paragraph in footer.paragraphs:
                        self._replace_in_paragraph(paragraph, replacement_map)
        
        # Remove images if requested
        if remove_images:
            self._remove_docx_images(doc)
        
        doc.save(output_path)
        return output_path
    
    def _replace_in_paragraph(self, paragraph, replacement_map: dict):
        """Replace text in a paragraph while preserving formatting."""
        if not paragraph.text:
            return
        
        # Get full paragraph text
        full_text = paragraph.text
        
        # Check if any replacements apply
        needs_replacement = False
        for original in replacement_map:
            if original in full_text:
                needs_replacement = True
                break
        
        if not needs_replacement:
            return
        
        # Apply all replacements
        new_text = full_text
        for original, anonymized in replacement_map.items():
            new_text = new_text.replace(original, anonymized)
        
        # If text changed, update runs
        if new_text != full_text:
            # Simple approach: clear and set new text
            # This may lose some formatting but is more reliable
            if paragraph.runs:
                # Store first run's formatting
                first_run = paragraph.runs[0]
                
                # Clear all runs
                for run in paragraph.runs[1:]:
                    run.clear()
                
                # Set new text in first run
                first_run.text = new_text
            else:
                paragraph.text = new_text
    
    def _remove_docx_images(self, doc):
        """Remove all images from a Word document."""
        # Access the document's XML
        for rel in list(doc.part.rels.values()):
            if "image" in rel.reltype:
                # Can't easily delete, but we can mark for removal
                pass
        
        # Alternative: iterate through elements and remove drawings
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                # Look for inline images
                drawing_elements = run._element.findall(
                    './/' + qn('w:drawing')
                )
                for drawing in drawing_elements:
                    drawing.getparent().remove(drawing)
    
    def write_xlsx(
        self, 
        original_path: Path, 
        output_path: Path,
        result: AnonymizationResult,
        highlight_changes: bool = False,
        anonymize_dates: bool = False
    ) -> Path:
        """
        Create an anonymized Excel workbook.
        
        Args:
            original_path: Path to original workbook
            output_path: Path to save anonymized workbook
            result: AnonymizationResult with replacements
            highlight_changes: Whether to highlight changed cells
            anonymize_dates: Whether dates setting is enabled
            
        Returns:
            Path to the saved workbook
        """
        wb = openpyxl.load_workbook(original_path)
        
        # Build replacement map, excluding date replacements for Excel cells
        replacement_map = {}
        for r in result.replacements:
            # Skip date-type replacements - they'll be handled separately for datetime cells
            if r.replacement_type != "date":
                replacement_map[r.original] = r.anonymized
        
        # Highlight fill for changed cells
        highlight_fill = PatternFill(
            start_color="FFFF00",
            end_color="FFFF00",
            fill_type="solid"
        ) if highlight_changes else None
        
        # Date counter for Excel datetime cells
        date_counter = 0
        
        # Process all sheets
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        # Check if cell contains datetime - handle specially
                        if isinstance(cell.value, (datetime, date)):
                            if anonymize_dates:
                                date_counter += 1
                                cell.value = f"Дата {date_counter}"
                                if highlight_fill:
                                    cell.fill = highlight_fill
                            # Skip further processing for datetime cells
                            continue
                        
                        # For non-datetime cells, apply text replacements
                        original_value = str(cell.value)
                        new_value = original_value
                        
                        # Apply replacements
                        for original, anonymized in replacement_map.items():
                            if original in new_value:
                                new_value = new_value.replace(original, anonymized)
                        
                        # Update cell if changed
                        if new_value != original_value:
                            # Try to preserve numeric type
                            if cell.data_type == 'n':
                                # Check if result is still numeric
                                clean_value = new_value.replace(' ', '').replace(',', '.')
                                if clean_value.replace('.', '').replace('-', '').isdigit():
                                    try:
                                        cell.value = float(clean_value) if '.' in clean_value else int(clean_value)
                                    except ValueError:
                                        cell.value = new_value
                                else:
                                    cell.value = new_value
                            else:
                                cell.value = new_value
                            
                            if highlight_fill:
                                cell.fill = highlight_fill
        
        wb.save(output_path)
        wb.close()
        return output_path
    
    def write_pdf(
        self, 
        original_path: Path, 
        output_path: Path,
        result: AnonymizationResult
    ) -> tuple[Path, list[str]]:
        """
        Create an anonymized PDF document.
        
        Args:
            original_path: Path to original PDF
            output_path: Path to save anonymized PDF
            result: AnonymizationResult with replacements
            
        Returns:
            Tuple of (path to saved PDF, list of warnings)
        """
        warnings = []
        
        if not HAS_PIKEPDF:
            import shutil
            shutil.copy(original_path, output_path)
            warnings.append("pikepdf не установлен - PDF скопирован без изменений")
            return output_path, warnings
        
        # Build replacement map
        replacement_map = {r.original: r.anonymized for r in result.replacements}
        
        if not replacement_map:
            import shutil
            shutil.copy(original_path, output_path)
            return output_path, warnings
        
        replacements_made = 0
        
        try:
            with pikepdf.open(original_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    # Try to modify content streams
                    try:
                        if "/Contents" in page:
                            contents = page["/Contents"]
                            
                            # Handle array of content streams
                            if isinstance(contents, pikepdf.Array):
                                for content_ref in contents:
                                    count = self._process_pdf_stream(content_ref, replacement_map)
                                    replacements_made += count
                            else:
                                # Single content stream
                                count = self._process_pdf_stream(contents, replacement_map)
                                replacements_made += count
                    except Exception:
                        pass  # Continue with other pages
                    
                    # Remove annotations
                    if "/Annots" in page:
                        try:
                            del page["/Annots"]
                        except Exception:
                            pass
                
                pdf.save(output_path)
                
        except Exception as e:
            import shutil
            shutil.copy(original_path, output_path)
            warnings.append(f"Ошибка обработки PDF: {str(e)}")
            return output_path, warnings
        
        if replacements_made == 0 and len(replacement_map) > 0:
            warnings.append("PDF: замены не применены. Текст может быть в изображениях или закодирован.")
        else:
            warnings.append(f"PDF: применено {replacements_made} замен в потоках")
        
        return output_path, warnings
    
    def _process_pdf_stream(self, stream_obj, replacement_map: dict) -> int:
        """Process a single PDF content stream."""
        count = 0
        
        try:
            if not hasattr(stream_obj, 'read_bytes'):
                return 0
            
            raw_content = stream_obj.read_bytes()
            
            # Try different decompression methods
            decompressed = None
            
            # Check if already decompressed
            try:
                decompressed = raw_content.decode('latin-1')
            except Exception:
                pass
            
            # Try zlib decompression (FlateDecode)
            if decompressed is None or not self._looks_like_text(decompressed):
                try:
                    decompressed = zlib.decompress(raw_content).decode('latin-1')
                except Exception:
                    pass
            
            if decompressed is None:
                return 0
            
            # Apply replacements in content stream
            modified = decompressed
            for original, anonymized in replacement_map.items():
                # Direct replacement
                if original in modified:
                    modified = modified.replace(original, anonymized)
                    count += 1
                
                # Try PDF string format replacements
                # Parenthesized strings: (text)
                if f"({original})" in modified:
                    modified = modified.replace(f"({original})", f"({anonymized})")
                    count += 1
                
                # Try hex-encoded strings
                hex_orig = original.encode('utf-16-be').hex().upper()
                hex_anon = anonymized.encode('utf-16-be').hex().upper()
                if hex_orig in modified:
                    modified = modified.replace(hex_orig, hex_anon)
                    count += 1
            
            if count > 0:
                # Re-compress and write back
                try:
                    compressed = zlib.compress(modified.encode('latin-1'))
                    stream_obj.write(compressed, filter=pikepdf.Name.FlateDecode)
                except Exception:
                    # Try writing uncompressed
                    stream_obj.write(modified.encode('latin-1'))
            
            return count
            
        except Exception:
            return 0
    
    def _looks_like_text(self, data: str) -> bool:
        """Check if data looks like readable text/commands."""
        # PDF content streams have operators like BT, ET, Tj, TJ
        pdf_operators = ['BT', 'ET', 'Tj', 'TJ', 'Tm', 'Td', 'cm', 'q', 'Q']
        return any(op in data for op in pdf_operators)
    
    def write_document(
        self, 
        original_path: Path, 
        output_path: Path,
        result: AnonymizationResult,
        file_type: str,
        settings: dict
    ) -> tuple[Path, list[str]]:
        """
        Write an anonymized document based on its type.
        
        Args:
            original_path: Path to original document
            output_path: Path to save anonymized document
            result: AnonymizationResult with replacements
            file_type: Type of document (docx, xlsx, pdf)
            settings: Anonymization settings
            
        Returns:
            Tuple of (path to saved document, list of warnings)
        """
        remove_images = settings.get("logos", False)
        anonymize_dates = settings.get("dates", False)
        warnings = []
        
        if file_type == "docx":
            path = self.write_docx(original_path, output_path, result, remove_images)
            return path, warnings
        elif file_type in ["xlsx", "xls"]:
            path = self.write_xlsx(original_path, output_path, result, 
                                   highlight_changes=False, anonymize_dates=anonymize_dates)
            return path, warnings
        elif file_type == "pdf":
            return self.write_pdf(original_path, output_path, result)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def create_comparison_preview(
        self, 
        original_text: str, 
        anonymized_text: str,
        max_length: int = 5000
    ) -> dict:
        """
        Create a comparison preview for the UI.
        
        Args:
            original_text: Original document text
            anonymized_text: Anonymized document text
            max_length: Maximum length for preview
            
        Returns:
            Dictionary with original and anonymized previews
        """
        # Truncate if too long
        original_preview = original_text[:max_length]
        anonymized_preview = anonymized_text[:max_length]
        
        if len(original_text) > max_length:
            original_preview += "\n\n... [текст обрезан для предпросмотра]"
        
        if len(anonymized_text) > max_length:
            anonymized_preview += "\n\n... [текст обрезан для предпросмотра]"
        
        return {
            "original": original_preview,
            "anonymized": anonymized_preview,
            "original_length": len(original_text),
            "anonymized_length": len(anonymized_text),
            "truncated": len(original_text) > max_length
        }


