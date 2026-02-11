"""Document parsing module for extracting content from various file formats."""

import io
from pathlib import Path
from typing import Any, Optional
from dataclasses import dataclass, field

# Document processing libraries
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
import openpyxl
from openpyxl.worksheet.worksheet import Worksheet
import pdfplumber
from PyPDF2 import PdfReader

# OCR support
try:
    from pdf2image import convert_from_path
    import pytesseract
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

from PIL import Image


@dataclass
class TextBlock:
    """Represents a block of text with its location."""
    text: str
    block_type: str  # 'paragraph', 'table_cell', 'header', 'footer'
    page: int = 0
    position: dict = field(default_factory=dict)


@dataclass
class ImageBlock:
    """Represents an image in the document."""
    image_data: bytes
    format: str
    block_id: str
    page: int = 0
    position: dict = field(default_factory=dict)


@dataclass
class ParsedDocument:
    """Container for parsed document content."""
    file_path: Path
    file_type: str
    text_blocks: list[TextBlock] = field(default_factory=list)
    images: list[ImageBlock] = field(default_factory=list)
    tables: list[list[list[str]]] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)
    raw_text: str = ""
    
    def get_full_text(self) -> str:
        """Get all text content as a single string."""
        if self.raw_text:
            return self.raw_text
        return "\n".join(block.text for block in self.text_blocks)


class DocumentParser:
    """Parser for extracting content from documents."""
    
    def __init__(self, ocr_language: str = "rus+eng", ocr_dpi: int = 300):
        self.ocr_language = ocr_language
        self.ocr_dpi = ocr_dpi
    
    def parse(self, file_path: Path) -> ParsedDocument:
        """
        Parse a document and extract its content.
        
        Args:
            file_path: Path to the document
            
        Returns:
            ParsedDocument with extracted content
        """
        suffix = file_path.suffix.lower()
        
        if suffix == ".docx":
            return self._parse_docx(file_path)
        elif suffix in [".xlsx", ".xls"]:
            return self._parse_xlsx(file_path)
        elif suffix == ".pdf":
            return self._parse_pdf(file_path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")
    
    def _parse_docx(self, file_path: Path) -> ParsedDocument:
        """Parse a Word document."""
        doc = Document(file_path)
        parsed = ParsedDocument(file_path=file_path, file_type="docx")
        
        # Extract metadata
        core_props = doc.core_properties
        parsed.metadata = {
            "author": core_props.author or "",
            "title": core_props.title or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "last_modified_by": core_props.last_modified_by or "",
            "revision": core_props.revision or 0,
            "category": core_props.category or "",
            "comments": core_props.comments or "",
        }
        
        # Extract paragraphs
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                parsed.text_blocks.append(TextBlock(
                    text=para.text,
                    block_type="paragraph",
                    position={"index": i}
                ))
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
                # Add cell text to text blocks
                for cell_idx, cell in enumerate(row.cells):
                    if cell.text.strip():
                        parsed.text_blocks.append(TextBlock(
                            text=cell.text,
                            block_type="table_cell",
                            position={"table": table_idx, "cell": cell_idx}
                        ))
            parsed.tables.append(table_data)
        
        # Extract images
        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.reltype:
                try:
                    image_part = rel.target_part
                    image_data = image_part.blob
                    image_format = image_part.content_type.split("/")[-1]
                    parsed.images.append(ImageBlock(
                        image_data=image_data,
                        format=image_format,
                        block_id=rel_id
                    ))
                except Exception:
                    pass
        
        # Compile raw text
        parsed.raw_text = parsed.get_full_text()
        
        return parsed
    
    def _parse_xlsx(self, file_path: Path) -> ParsedDocument:
        """Parse an Excel spreadsheet."""
        wb = openpyxl.load_workbook(file_path, data_only=True)
        parsed = ParsedDocument(file_path=file_path, file_type="xlsx")
        
        # Extract metadata
        props = wb.properties
        parsed.metadata = {
            "creator": props.creator or "",
            "title": props.title or "",
            "subject": props.subject or "",
            "keywords": props.keywords or "",
            "created": str(props.created) if props.created else "",
            "modified": str(props.modified) if props.modified else "",
            "lastModifiedBy": props.lastModifiedBy or "",
            "category": props.category or "",
            "description": props.description or "",
        }
        
        # Extract content from all sheets
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_data = []
            
            for row_idx, row in enumerate(ws.iter_rows(values_only=True)):
                row_data = []
                for col_idx, cell_value in enumerate(row):
                    if cell_value is not None:
                        cell_text = str(cell_value)
                        row_data.append(cell_text)
                        parsed.text_blocks.append(TextBlock(
                            text=cell_text,
                            block_type="table_cell",
                            position={
                                "sheet": sheet_name,
                                "row": row_idx,
                                "col": col_idx
                            }
                        ))
                    else:
                        row_data.append("")
                if any(row_data):
                    sheet_data.append(row_data)
            
            if sheet_data:
                parsed.tables.append(sheet_data)
        
        # Extract images from sheets
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            for idx, image in enumerate(ws._images):
                try:
                    img_data = image._data()
                    parsed.images.append(ImageBlock(
                        image_data=img_data,
                        format="png",
                        block_id=f"{sheet_name}_img_{idx}"
                    ))
                except Exception:
                    pass
        
        wb.close()
        parsed.raw_text = parsed.get_full_text()
        
        return parsed
    
    def _parse_pdf(self, file_path: Path) -> ParsedDocument:
        """Parse a PDF document."""
        parsed = ParsedDocument(file_path=file_path, file_type="pdf")
        
        # Try to extract text directly first
        try:
            with pdfplumber.open(file_path) as pdf:
                # Extract metadata
                parsed.metadata = pdf.metadata or {}
                
                all_text = []
                for page_num, page in enumerate(pdf.pages):
                    # Extract text
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        parsed.text_blocks.append(TextBlock(
                            text=page_text,
                            block_type="paragraph",
                            page=page_num
                        ))
                        all_text.append(page_text)
                    
                    # Extract tables
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            parsed.tables.append(table)
                    
                    # Extract images
                    for img_idx, img in enumerate(page.images):
                        try:
                            # Get image data if available
                            parsed.images.append(ImageBlock(
                                image_data=b"",  # Placeholder
                                format="unknown",
                                block_id=f"page_{page_num}_img_{img_idx}",
                                page=page_num,
                                position={"x0": img["x0"], "y0": img["top"]}
                            ))
                        except Exception:
                            pass
                
                parsed.raw_text = "\n".join(all_text)
        
        except Exception as e:
            print(f"Error parsing PDF with pdfplumber: {e}")
        
        # If no text extracted, try OCR
        if not parsed.raw_text.strip() and HAS_OCR:
            parsed = self._ocr_pdf(file_path, parsed)
        
        # Try PyPDF2 for additional metadata
        try:
            reader = PdfReader(file_path)
            if reader.metadata:
                parsed.metadata.update({
                    k.replace("/", ""): v 
                    for k, v in reader.metadata.items() 
                    if v
                })
        except Exception:
            pass
        
        return parsed
    
    def _ocr_pdf(self, file_path: Path, parsed: ParsedDocument) -> ParsedDocument:
        """Perform OCR on a scanned PDF."""
        if not HAS_OCR:
            return parsed
        
        try:
            images = convert_from_path(file_path, dpi=self.ocr_dpi)
            
            all_text = []
            for page_num, image in enumerate(images):
                # Perform OCR
                text = pytesseract.image_to_string(image, lang=self.ocr_language)
                if text.strip():
                    parsed.text_blocks.append(TextBlock(
                        text=text,
                        block_type="ocr_text",
                        page=page_num
                    ))
                    all_text.append(text)
                
                # Store page image for potential logo detection
                img_buffer = io.BytesIO()
                image.save(img_buffer, format="PNG")
                parsed.images.append(ImageBlock(
                    image_data=img_buffer.getvalue(),
                    format="png",
                    block_id=f"ocr_page_{page_num}",
                    page=page_num
                ))
            
            parsed.raw_text = "\n".join(all_text)
        
        except Exception as e:
            print(f"OCR error: {e}")
        
        return parsed
    
    def is_scanned_pdf(self, file_path: Path) -> bool:
        """Check if a PDF is scanned (image-based) rather than text-based."""
        try:
            with pdfplumber.open(file_path) as pdf:
                # Check first few pages
                for page in pdf.pages[:3]:
                    text = page.extract_text() or ""
                    if len(text.strip()) > 50:  # Has substantial text
                        return False
                return True
        except Exception:
            return False


