from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel
from datetime import datetime
from typing import Optional, List
import uuid
import os
import hashlib
import io

from database import get_db
from models.document import Document, DocumentVersion, DocumentStatus
from models.user import User
from services.document_processor import DocumentProcessor
from services.auth_service import get_current_user
from config import settings

router = APIRouter()

# Response models
class DocumentResponse(BaseModel):
    id: str
    name: str
    description: Optional[str] = None
    original_filename: str
    file_type: Optional[str] = None
    file_size: Optional[int] = None
    page_count: Optional[int] = None
    status: str
    folder: Optional[str] = None
    uploaded_at: datetime
    version_count: int = 1

class DocumentVersionResponse(BaseModel):
    id: str
    version_number: int
    created_at: datetime
    change_summary: Optional[str] = None
    change_count: int = 0
    critical_changes: int = 0
    major_changes: int = 0
    minor_changes: int = 0

class TimelineResponse(BaseModel):
    document_id: str
    document_name: str
    versions: List[DocumentVersionResponse]

class DocumentListResponse(BaseModel):
    documents: List[DocumentResponse]
    total: int
    page: int
    page_size: int


@router.post("/upload", response_model=DocumentResponse)
async def upload_document(
    file: UploadFile = File(...),
    name: Optional[str] = None,
    description: Optional[str] = None,
    folder: Optional[str] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Upload a new document (auth required)"""
    # Validate file extension
    ext = file.filename.split(".")[-1].lower() if "." in file.filename else ""
    if ext not in settings.ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"File type .{ext} not allowed. Allowed: {settings.ALLOWED_EXTENSIONS}")
    
    # Read file content
    content = await file.read()
    if len(content) > settings.MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File too large")
    
    # Generate file path
    file_id = str(uuid.uuid4())
    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
    file_path = os.path.join(settings.UPLOAD_DIR, f"{file_id}.{ext}")
    
    # Save file
    with open(file_path, "wb") as f:
        f.write(content)
    
    # Calculate hash
    content_hash = hashlib.sha256(content).hexdigest()
    
    # Process document
    processor = DocumentProcessor()
    extracted_text, page_count = processor.extract_text(file_path, ext)
    
    # Create document with user ownership
    doc = Document(
        id=file_id,
        tenant_id=current_user.tenant_id or "default",
        name=name or file.filename,
        description=description,
        file_path=file_path,
        original_filename=file.filename,
        file_size=len(content),
        page_count=page_count,
        uploaded_by=current_user.id,  # Set document owner
        uploaded_at=datetime.utcnow(),
        status=DocumentStatus.READY.value,
        content_hash=content_hash,
        extracted_text=extracted_text,
        folder=folder
    )
    db.add(doc)
    
    # Create initial version
    version = DocumentVersion(
        id=str(uuid.uuid4()),
        document_id=file_id,
        version_number=1,
        content=extracted_text,
        file_path=file_path,
        created_by=current_user.id,
        created_at=datetime.utcnow()
    )
    db.add(version)
    
    db.commit()
    db.refresh(doc)
    
    return DocumentResponse(
        id=doc.id,
        name=doc.name,
        description=doc.description,
        original_filename=doc.original_filename,
        file_type=ext,
        file_size=doc.file_size,
        page_count=doc.page_count,
        status=doc.status,
        folder=doc.folder,
        uploaded_at=doc.uploaded_at,
        version_count=1
    )

@router.get("/", response_model=DocumentListResponse)
async def list_documents(
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    folder: Optional[str] = None,
    search: Optional[str] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """List documents for current user (auth required)"""
    # Filter by current user's documents only
    query = db.query(Document).filter(
        Document.is_archived == "false",
        Document.uploaded_by == current_user.id  # Only user's own documents
    )
    
    # Filter by folder
    if folder:
        query = query.filter(Document.folder == folder)
    
    # Search
    if search:
        query = query.filter(
            (Document.name.ilike(f"%{search}%")) |
            (Document.extracted_text.ilike(f"%{search}%"))
        )
    
    total = query.count()
    documents = query.order_by(Document.uploaded_at.desc()).offset((page - 1) * page_size).limit(page_size).all()
    
    return DocumentListResponse(
        documents=[
            DocumentResponse(
                id=doc.id,
                name=doc.name,
                description=doc.description,
                original_filename=doc.original_filename,
                file_type=doc.original_filename.split(".")[-1] if "." in doc.original_filename else None,
                file_size=doc.file_size,
                page_count=doc.page_count,
                status=doc.status,
                folder=doc.folder,
                uploaded_at=doc.uploaded_at,
                version_count=len(doc.versions) if doc.versions else 1
            )
            for doc in documents
        ],
        total=total,
        page=page,
        page_size=page_size
    )

@router.get("/{document_id}", response_model=DocumentResponse)
async def get_document(
    document_id: str, 
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get document by ID (auth required, owner only)"""
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.uploaded_by == current_user.id
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return DocumentResponse(
        id=doc.id,
        name=doc.name,
        description=doc.description,
        original_filename=doc.original_filename,
        file_type=doc.original_filename.split(".")[-1] if "." in doc.original_filename else None,
        file_size=doc.file_size,
        page_count=doc.page_count,
        status=doc.status,
        folder=doc.folder,
        uploaded_at=doc.uploaded_at,
        version_count=len(doc.versions) if doc.versions else 1
    )

@router.delete("/{document_id}")
async def delete_document(
    document_id: str, 
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Delete (archive) document (auth required, owner only)"""
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.uploaded_by == current_user.id
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    doc.is_archived = "true"
    db.commit()
    
    return {"message": "Document archived", "id": document_id}

@router.get("/{document_id}/versions", response_model=List[DocumentVersionResponse])
async def get_versions(
    document_id: str, 
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get all versions of a document (auth required, owner only)"""
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.uploaded_by == current_user.id
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return [
        DocumentVersionResponse(
            id=v.id,
            version_number=v.version_number,
            created_at=v.created_at,
            change_summary=v.change_summary,
            change_count=v.change_count or 0,
            critical_changes=v.critical_changes or 0,
            major_changes=v.major_changes or 0,
            minor_changes=v.minor_changes or 0
        )
        for v in (doc.versions or [])
    ]

@router.get("/{document_id}/timeline", response_model=TimelineResponse)
async def get_timeline(
    document_id: str, 
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get timeline visualization data for a document (auth required, owner only)"""
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.uploaded_by == current_user.id
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return TimelineResponse(
        document_id=doc.id,
        document_name=doc.name,
        versions=[
            DocumentVersionResponse(
                id=v.id,
                version_number=v.version_number,
                created_at=v.created_at,
                change_summary=v.change_summary,
                change_count=v.change_count or 0,
                critical_changes=v.critical_changes or 0,
                major_changes=v.major_changes or 0,
                minor_changes=v.minor_changes or 0
            )
            for v in (doc.versions or [])
        ]
    )

@router.get("/{document_id}/content")
async def get_document_content(
    document_id: str, 
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get extracted text content of a document (auth required, owner only)"""
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.uploaded_by == current_user.id
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return {
        "id": doc.id,
        "name": doc.name,
        "content": doc.extracted_text,
        "page_count": doc.page_count
    }


@router.get("/{document_id}/download")
async def download_document(
    document_id: str, 
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Download document as DOCX file (auth required, owner only)"""
    from urllib.parse import quote
    
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.uploaded_by == current_user.id
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Helper to create proper Content-Disposition header with UTF-8 support
    def get_content_disposition(filename: str) -> str:
        # RFC 5987 encoding for non-ASCII filenames
        encoded_filename = quote(filename, safe='')
        return f"attachment; filename*=UTF-8''{encoded_filename}"
    
    # If original file exists and is a docx, return it
    if doc.file_path and os.path.exists(doc.file_path) and doc.file_path.endswith('.docx'):
        def iterfile():
            with open(doc.file_path, "rb") as f:
                yield from f
        
        filename = doc.name if doc.name.endswith('.docx') else f"{doc.name}.docx"
        return StreamingResponse(
            iterfile(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": get_content_disposition(filename)}
        )
    
    # Otherwise, create a DOCX from extracted text
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        docx_doc = DocxDocument()
        
        # Set document properties
        docx_doc.core_properties.title = doc.name
        docx_doc.core_properties.author = "СравнениеДок Платформа"
        
        # Add content
        content = doc.extracted_text or ""
        
        # Split content by paragraphs and add them
        paragraphs = content.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                p = docx_doc.add_paragraph(para_text.strip())
                p.style.font.name = 'Times New Roman'
                p.style.font.size = Pt(12)
        
        # Save to bytes buffer
        buffer = io.BytesIO()
        docx_doc.save(buffer)
        buffer.seek(0)
        
        filename = doc.name if doc.name.endswith('.docx') else f"{doc.name}.docx"
        
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": get_content_disposition(filename)}
        )
        
    except ImportError:
        # If python-docx not available, return as plain text
        content = doc.extracted_text or ""
        buffer = io.BytesIO(content.encode('utf-8'))
        
        filename = doc.name if doc.name.endswith('.txt') else f"{doc.name}.txt"
        
        return StreamingResponse(
            buffer,
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": get_content_disposition(filename)}
        )
