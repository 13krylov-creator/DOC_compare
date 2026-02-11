"""FastAPI router for Document Anonymizer - integrated into main platform."""

import json
import asyncio
import re
from pathlib import Path
from typing import Optional
from datetime import datetime
from uuid import uuid4

from fastapi import APIRouter, File, UploadFile, Form, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse, JSONResponse, HTMLResponse

from config import (
    ANONYMIZER_UPLOAD_DIR,
    ANONYMIZER_SUPPORTED_FORMATS,
    ANONYMIZATION_OPTIONS,
    DEFAULT_PROFILES,
    ML_CONFIG,
    settings,
)
from anonymizer_utils.file_utils import (
    save_uploaded_file,
    get_file_extension,
    get_output_path,
    get_task_files,
    cleanup_old_files,
    delete_task_files,
)
from anonymizer_core.document_parser import DocumentParser
from anonymizer_core.anonymizer import Anonymizer, Replacement
from anonymizer_core.document_writer import DocumentWriter
from anonymizer_core.metadata_cleaner import MetadataCleaner
from anonymizer_core.ml_integration import MLIntegration
from anonymizer_core.validator import Validator


router = APIRouter()

# In-memory task storage
tasks = {}

# Initialize components
parser = DocumentParser()
ml_integration = MLIntegration()
anonymizer = Anonymizer(ml_integration)
writer = DocumentWriter()
metadata_cleaner = MetadataCleaner()
validator = Validator(ml_integration)

MAX_FILE_SIZE_BYTES = settings.MAX_FILE_SIZE


# ============ Helper Functions ============

def add_log(task: dict, message: str):
    """Add a log entry to task."""
    if "logs" not in task:
        task["logs"] = []
    task["logs"].append({
        "time": datetime.now().strftime("%H:%M:%S"),
        "message": message
    })
    task["message"] = message


async def add_log_async(task: dict, message: str):
    """Add a log entry and yield control to event loop."""
    add_log(task, message)
    await asyncio.sleep(0)


def save_as_markdown(text: str, output_path: Path, original_name: str) -> tuple:
    """Save text as markdown file."""
    header = f"""# Обезличенный документ

**Источник:** {original_name}

---

"""
    content = header + text
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)
    return output_path, []


def create_word_from_text(text: str, output_path: Path, original_name: str) -> tuple:
    """Create a Word document from plain text."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("Обезличенный документ", level=1)
    doc.add_paragraph(f"Источник: {original_name}")
    doc.add_paragraph("---")

    paragraphs = text.split("\n")
    for para_text in paragraphs:
        if para_text.strip():
            p = doc.add_paragraph(para_text)
            p.style.font.size = Pt(11)

    doc.save(output_path)
    return output_path, []


async def process_pdf_with_chandra_async(pdf_path: Path, ml_int, task) -> tuple:
    """Async version: Process PDF using Chandra OCR to get markdown with tables."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        await add_log_async(task, "[!] PyMuPDF not installed")
        return "", []

    logs = []
    all_markdown = []

    doc = fitz.open(pdf_path)
    total_pages = len(doc)
    await add_log_async(task, f"PDF: {total_pages} страниц")

    for page_num in range(total_pages):
        await add_log_async(task, f"[OCR] Страница {page_num + 1}/{total_pages}...")
        task["progress"] = 10 + int((page_num / total_pages) * 10)

        page = doc.load_page(page_num)
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        img_bytes = pix.tobytes("png")

        markdown = ml_int.ocr_with_chandra(img_bytes)

        if markdown and len(markdown) > 10:
            all_markdown.append(f"## Страница {page_num + 1}\n\n{markdown}")
            await add_log_async(task, f"[OK] Стр.{page_num + 1}: {len(markdown)} символов")
        else:
            await add_log_async(task, f"[!] Стр.{page_num + 1}: OCR не дал результата")

    doc.close()

    full_markdown = "\n\n---\n\n".join(all_markdown)
    return full_markdown, logs


async def process_document(task_id: str):
    """Background task for document processing."""
    task = tasks.get(task_id)
    if not task:
        return

    processing_warnings = []
    task["logs"] = []

    try:
        # Step 1: Check ML availability
        task["progress"] = 5
        await add_log_async(task, "[1/7] Проверка доступности ML моделей...")

        ml_status = ml_integration.is_available()
        if ml_status.get("gpt"):
            await add_log_async(task, "[OK] GPT модель доступна")
        else:
            await add_log_async(task, "[!] GPT модель недоступна - только regex")
            processing_warnings.append("GPT недоступна - используются только regex-паттерны")

        if ml_status.get("vision"):
            await add_log_async(task, "[OK] Vision модель доступна")
        else:
            await add_log_async(task, "[!] Vision модель недоступна")

        # Step 2: Parse document
        task["progress"] = 10
        await add_log_async(task, f"[2/7] Загрузка файла: {task['filename']}")

        original_path = Path(task["original_path"])
        file_type = task["file_type"]

        # For PDF, use Chandra OCR with markdown
        if file_type == "pdf" and ml_status.get("vision"):
            await add_log_async(task, "PDF обнаружен - используем Chandra OCR...")
            markdown_text, ocr_logs = await process_pdf_with_chandra_async(original_path, ml_integration, task)
            parsed_raw_text = markdown_text
            task["is_markdown"] = True
        else:
            await add_log_async(task, f"Извлечение текста из {file_type.upper()}...")
            parsed = parser.parse(original_path)
            parsed_raw_text = parsed.raw_text
            task["is_markdown"] = False

        text_length = len(parsed_raw_text)
        await add_log_async(task, f"Извлечено {text_length} символов текста")

        if not parsed_raw_text.strip():
            await add_log_async(task, "[!] Текст не найден")
            processing_warnings.append("Не удалось извлечь текст из документа")

        # Step 3: ML-based entity detection
        task["progress"] = 25
        ml_entities = {"companies": [], "persons": [], "prices": []}

        if ml_status.get("gpt") and parsed_raw_text.strip():
            await add_log_async(task, "[3/7] Анализ текста через GPT (постранично)...")

            pages = re.split(r'---\s*\n\s*## Страница', parsed_raw_text)
            if len(pages) == 1:
                chunk_size = 3000
                pages = [parsed_raw_text[i:i+chunk_size] for i in range(0, len(parsed_raw_text), chunk_size)]

            await add_log_async(task, f"Разбито на {len(pages)} частей")

            for i, page_text in enumerate(pages[:10]):
                if len(page_text.strip()) < 50:
                    continue

                await add_log_async(task, f"[GPT] Часть {i+1}/{min(len(pages), 10)}...")

                clean_text = re.sub(r'<[^>]+>', ' ', page_text)
                clean_text = re.sub(r'\s+', ' ', clean_text)[:3000]

                try:
                    prompt = f"""You are an information extraction assistant for Russian documents.
Extract ONLY explicitly present data. Do NOT invent anything.

CATEGORIES TO EXTRACT:
1. company_names: Legal entity names with ООО/АО/ПАО/ЗАО/ИП prefix, OR well-known brands (e.g. "НИП-центр", "НИР-центр", "Уралмеханобр", "УГМК")
2. person_contacts: Full names (Фамилия Имя Отчество) or names with initials (Иванов И.П.)
3. prices_amounts: Monetary values with currency (1 500 000 руб., 45 000,00 ₽, итого: 250 тыс.)

STRICT EXCLUSIONS (NEVER include):
- Column headers: Наименование, Сумма, Итого, ИНН, Номер, Дата, Стоимость, Адрес
- Document types: Счет-фактура, Акт, Договор, Платежное поручение, Реестр
- Roles: Продавец, Покупатель, Получатель, Заказчик, Исполнитель
- Work types: СМР, ПИР, Материалы, ТМЦ, Оборудование
- Terms: Объект, Стройка, Капитальные вложения, Инфраструктура
- Anything in quotes «» without ООО/АО/ПАО prefix
- Code patterns: ЮМ 2024, ВГОК 2025, Реестр КВиК (these are codes, NOT companies)
- Generic phrases containing: указывается, примечание, проверка, сверка

VALIDATION RULES:
- company_names MUST have legal form (ООО/АО/ПАО/ЗАО/ИП) OR be a recognizable brand name
- person_contacts MUST look like actual human names (Фамилия + Имя or initials)
- prices_amounts MUST have numeric value AND currency indicator

Return ONLY valid JSON (no markdown, no explanation):
{{"company_names":[],"person_contacts":[],"prices_amounts":[]}}

TEXT TO ANALYZE:
{clean_text}"""

                    gpt_response, gpt_error = ml_integration.ask_gpt(prompt, max_retries=3)

                    if gpt_response:
                        json_match = re.search(r'\{[^{}]*\}', gpt_response, re.DOTALL)
                        if json_match:
                            page_entities = json.loads(json_match.group())

                            from anonymizer_utils.stopwords import filter_gpt_results

                            category_map = {
                                "company_names": "companies",
                                "person_contacts": "persons",
                                "prices_amounts": "prices",
                                "companies": "companies",
                                "persons": "persons",
                                "prices": "prices"
                            }

                            companies_before = len(ml_entities["companies"])
                            persons_before = len(ml_entities["persons"])

                            for gpt_key, ml_key in category_map.items():
                                raw_items = page_entities.get(gpt_key, [])
                                if raw_items:
                                    await add_log_async(task, f"  GPT: {len(raw_items)} {gpt_key}")
                                filtered_items = filter_gpt_results(raw_items)
                                if len(filtered_items) < len(raw_items):
                                    await add_log_async(task, f"  Фильтр: {len(raw_items)} -> {len(filtered_items)}")
                                for item in filtered_items:
                                    if item and item not in ml_entities[ml_key]:
                                        ml_entities[ml_key].append(item)

                            companies_added = len(ml_entities["companies"]) - companies_before
                            persons_added = len(ml_entities["persons"]) - persons_before
                            await add_log_async(task, f"[OK] +{companies_added} компаний, +{persons_added} ФИО")
                        else:
                            await add_log_async(task, f"[!] Часть {i+1}: неверный JSON")
                    elif gpt_error:
                        await add_log_async(task, f"[!] Часть {i+1}: {gpt_error}")
                    else:
                        await add_log_async(task, f"[!] Часть {i+1}: GPT не ответил")

                except Exception as e:
                    await add_log_async(task, f"[!] Часть {i+1}: {str(e)[:40]}")

            total_found = len(ml_entities["companies"]) + len(ml_entities["persons"]) + len(ml_entities["prices"])
            await add_log_async(task, f"GPT всего: {len(ml_entities['companies'])} компаний, {len(ml_entities['persons'])} ФИО, {len(ml_entities['prices'])} сумм")

        # Step 4: Anonymization
        task["progress"] = 40
        await add_log_async(task, "[4/7] Начинаем обезличивание...")

        anonymizer.reset_counters()

        result = anonymizer.anonymize_text(
            parsed_raw_text,
            task["settings"],
            use_ml=False
        )

        # Apply ML-detected entities
        if ml_entities:
            if task["settings"].get("companies", False):
                for company in ml_entities.get("companies", []):
                    if company and len(company) > 2:
                        if company in result.anonymized_text:
                            anonymizer._company_counter += 1
                            replacement = f"Компания {anonymizer._company_counter}"
                            result.anonymized_text = result.anonymized_text.replace(company, replacement)
                            result.replacements.append(Replacement(
                                original=company,
                                anonymized=replacement,
                                replacement_type="company_gpt"
                            ))
                            await add_log_async(task, f"[+] Компания: {company[:25]}")

                        variations = [
                            company.upper(),
                            company.lower(),
                            company.replace("-", ""),
                            "НИ" + company if company.startswith("ИР") else None,
                            "НИИ" + company[1:] if company.startswith("И") else None,
                        ]
                        if "НИР" in company.upper():
                            variations.extend(["НИИР-ЦЕНТР", "НИР-ЦЕНТР", "НИР-Центр", "НИИР-центр"])

                        for var in variations:
                            if var and var in result.anonymized_text and var != company:
                                result.anonymized_text = result.anonymized_text.replace(var, f"Компания {anonymizer._company_counter}")
                                await add_log_async(task, f"[+] Вариация: {var[:25]}")

            if task["settings"].get("personal_data", False) or task["settings"].get("personal", False):
                for person in ml_entities.get("persons", []):
                    if person and len(person) > 2 and person in result.anonymized_text:
                        anonymizer._person_counter += 1
                        replacement = f"Контактное лицо {anonymizer._person_counter}"
                        result.anonymized_text = result.anonymized_text.replace(person, replacement)
                        result.replacements.append(Replacement(
                            original=person,
                            anonymized=replacement,
                            replacement_type="person_gpt"
                        ))
                        await add_log_async(task, f"[+] ФИО: {person[:25]}")

            if task["settings"].get("prices", False):
                for price in ml_entities.get("prices", []):
                    if price and len(price) > 3 and price in result.anonymized_text:
                        anonymizer._price_counter += 1
                        replacement = f"[СУММА {anonymizer._price_counter}]"
                        result.anonymized_text = result.anonymized_text.replace(price, replacement)
                        result.replacements.append(Replacement(
                            original=price,
                            anonymized=replacement,
                            replacement_type="price_gpt"
                        ))
                        await add_log_async(task, f"[+] Сумма: {price[:20]}")

        # Additional pass: catch remaining prices with enhanced regex
        if task["settings"].get("prices", False):
            price_patterns = [
                r'\d{1,3}(?:[\s\xa0]?\d{3})+(?:[,\.]\d{2})?\s*(?:руб(?:л(?:ей|я|ь)?)?\.?|₽)',
                r'\d+[,\.]\d{2}\s*(?:руб(?:л(?:ей|я|ь)?)?\.?|₽)',
                r'\d{1,3}(?:[\s\xa0]\d{3}){2,}(?:[,\.]\d{2})?',
            ]

            price_map = {}
            for pattern in price_patterns:
                for match in re.finditer(pattern, result.anonymized_text):
                    price_text = match.group()
                    if price_text in price_map:
                        continue
                    if price_text in [r.original for r in result.replacements]:
                        continue
                    if price_text.startswith('[СУММА'):
                        continue
                    clean_num = re.sub(r'[^\d]', '', price_text)
                    if len(clean_num) == 4 and not any(c in price_text for c in ['руб', '₽', 'рубл']):
                        continue
                    if len(clean_num) == 6 and not any(c in price_text for c in ['руб', '₽', 'рубл', ' ']):
                        continue

                    anonymizer._price_counter += 1
                    replacement_text = f"[СУММА {anonymizer._price_counter}]"
                    price_map[price_text] = replacement_text
                    result.replacements.append(Replacement(
                        original=price_text,
                        anonymized=replacement_text,
                        replacement_type="price_regex"
                    ))

            for price_text, replacement_text in price_map.items():
                result.anonymized_text = result.anonymized_text.replace(price_text, replacement_text)

        # Count by category
        categories = {}
        for r in result.replacements:
            cat = r.replacement_type
            categories[cat] = categories.get(cat, 0) + 1

        replacements_count = len(result.replacements)

        await add_log_async(task, f"=== Статистика обезличивания ===")
        await add_log_async(task, f"Всего замен: {replacements_count}")

        if categories.get("price", 0) > 0 or categories.get("price_regex", 0) > 0 or categories.get("price_gpt", 0) > 0:
            total_prices = categories.get("price", 0) + categories.get("price_regex", 0) + categories.get("price_gpt", 0)
            await add_log_async(task, f"├─ Цены и суммы: {total_prices}")

        if categories.get("company", 0) > 0 or categories.get("company_gpt", 0) > 0:
            total_companies = categories.get("company", 0) + categories.get("company_gpt", 0)
            await add_log_async(task, f"├─ Компании: {total_companies}")

        if categories.get("personal", 0) > 0 or categories.get("person_gpt", 0) > 0:
            total_persons = categories.get("personal", 0) + categories.get("person_gpt", 0)
            await add_log_async(task, f"├─ ФИО: {total_persons}")

        if categories.get("date", 0) > 0:
            await add_log_async(task, f"├─ Даты: {categories.get('date', 0)}")

        if categories.get("address", 0) > 0:
            await add_log_async(task, f"├─ Адреса: {categories.get('address', 0)}")

        if categories.get("requisites", 0) > 0:
            await add_log_async(task, f"├─ Реквизиты: {categories.get('requisites', 0)}")

        if categories.get("email", 0) > 0:
            await add_log_async(task, f"├─ Email: {categories.get('email', 0)}")

        if categories.get("phone", 0) > 0:
            await add_log_async(task, f"└─ Телефоны: {categories.get('phone', 0)}")

        if replacements_count == 0:
            await add_log_async(task, "[!] Данных для обезличивания не найдено")
            processing_warnings.append("Не найдено данных для обезличивания")

        # Step 5: Create output document
        task["progress"] = 60
        await add_log_async(task, "[5/7] Создание обезличенного документа...")

        file_type = task["file_type"]
        if file_type == "pdf" and task.get("is_markdown"):
            await add_log_async(task, "Сохранение как Markdown (таблицы сохранены)")
            output_filename = task["filename"].replace(".pdf", ".md")
            output_path = get_output_path(task_id, output_filename)
            output_path, write_warnings = save_as_markdown(
                result.anonymized_text,
                output_path,
                task["filename"]
            )
            task["output_ext"] = ".md"
        elif file_type == "pdf":
            await add_log_async(task, "PDF будет сохранен как Word")
            output_filename = task["filename"].replace(".pdf", ".docx")
            output_path = get_output_path(task_id, output_filename)
            output_path, write_warnings = create_word_from_text(
                result.anonymized_text,
                output_path,
                task["filename"]
            )
            task["output_ext"] = ".docx"
        else:
            output_path = get_output_path(task_id, task["filename"])
            output_path, write_warnings = writer.write_document(
                original_path,
                output_path,
                result,
                file_type,
                task["settings"]
            )
            task["output_ext"] = Path(task["filename"]).suffix

        processing_warnings.extend(write_warnings)
        await add_log_async(task, f"Файл сохранен: {output_path.name}")

        # Step 6: Clean metadata
        task["progress"] = 80
        output_ext = task.get("output_ext", "")
        if task["settings"].get("metadata", False) and output_ext not in (".md", ".txt"):
            await add_log_async(task, "[6/7] Очистка метаданных...")
            cleaned_path = output_path.parent / f"cleaned_{output_path.name}"
            try:
                final_type = "docx" if file_type == "pdf" else file_type
                metadata_cleaner.clean_document(output_path, cleaned_path, final_type)
                cleaned_path.replace(output_path)
                await add_log_async(task, "[OK] Метаданные очищены")
            except Exception as e:
                await add_log_async(task, f"[!] Ошибка очистки: {str(e)[:30]}")
                processing_warnings.append(f"Ошибка очистки метаданных: {str(e)}")
        elif task["settings"].get("metadata", False):
            await add_log_async(task, "[6/7] Пропуск очистки метаданных (текстовый формат)")

        # Step 7: Validation
        task["progress"] = 90
        await add_log_async(task, "[7/7] Валидация результата...")

        validation_result = validator.validate_regex(result.anonymized_text, task["settings"])

        all_warnings = processing_warnings + validation_result.warnings

        # Done!
        task["progress"] = 100
        task["status"] = "done"
        await add_log_async(task, "[OK] Обработка завершена успешно!")

        task["output_path"] = str(output_path)
        task["original_text"] = parsed_raw_text
        task["anonymized_text"] = result.anonymized_text
        task["replacements"] = [
            {"original": r.original, "anonymized": r.anonymized, "type": r.replacement_type}
            for r in result.replacements
        ]
        task["validation"] = {
            "is_valid": validation_result.is_valid,
            "issues": validation_result.issues,
            "warnings": all_warnings,
            "confidence": validation_result.confidence
        }

    except Exception as e:
        task["status"] = "error"
        add_log(task, f"[ERROR] Ошибка: {str(e)}")
        task["message"] = f"Ошибка обработки: {str(e)}"
        task["progress"] = 0


def markdown_to_html(md_text: str) -> str:
    """Convert markdown to HTML with basic formatting."""
    html = md_text

    html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
    html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
    html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)

    html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)
    html = re.sub(r'__(.+?)__', r'<strong>\1</strong>', html)
    html = re.sub(r'\*(.+?)\*', r'<em>\1</em>', html)

    html = re.sub(r'^---+$', r'<hr>', html, flags=re.MULTILINE)
    html = re.sub(r'^- (.+)$', r'<li>\1</li>', html, flags=re.MULTILINE)
    html = re.sub(r'^(\d+)\. (.+)$', r'<li>\2</li>', html, flags=re.MULTILINE)

    html = re.sub(r'\[СУММА (\d+)\]', r'<span class="anonymized">[СУММА \1]</span>', html)
    html = re.sub(r'Компания (\d+)', r'<span class="anonymized">Компания \1</span>', html)
    html = re.sub(r'Контактное лицо (\d+)', r'<span class="anonymized">Контактное лицо \1</span>', html)
    html = re.sub(r'Дата (\d+)', r'<span class="anonymized">Дата \1</span>', html)

    lines = html.split('\n')
    result_lines = []
    for line in lines:
        line = line.strip()
        if line and not line.startswith('<') and not line.startswith('|'):
            line = f'<p>{line}</p>'
        result_lines.append(line)
    html = '\n'.join(result_lines)

    html = convert_md_tables(html)
    return html


def convert_md_tables(text: str) -> str:
    """Convert markdown tables to HTML."""
    lines = text.split('\n')
    result_lines = []
    in_table = False
    table_rows = []

    for line in lines:
        if '|' in line and not line.strip().startswith('<'):
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if cells:
                if all(c.replace('-', '').replace(':', '') == '' for c in cells):
                    continue
                if not in_table:
                    in_table = True
                    table_rows = []
                table_rows.append(cells)
        else:
            if in_table:
                result_lines.append(render_table(table_rows))
                in_table = False
                table_rows = []
            result_lines.append(line)

    if in_table:
        result_lines.append(render_table(table_rows))

    return '\n'.join(result_lines)


def render_table(rows: list) -> str:
    """Render table rows as HTML."""
    if not rows:
        return ''

    html = '<table>\n'
    if rows:
        html += '<thead><tr>'
        for cell in rows[0]:
            html += f'<th>{cell}</th>'
        html += '</tr></thead>\n'

    if len(rows) > 1:
        html += '<tbody>\n'
        for row in rows[1:]:
            html += '<tr>'
            for cell in row:
                html += f'<td>{cell}</td>'
            html += '</tr>\n'
        html += '</tbody>\n'

    html += '</table>'
    return html


def create_pdf_from_html(html_content: str, output_path: Path, source_name: str):
    """Create PDF from markdown/HTML content using reportlab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import os

    font_name = 'Helvetica'
    try:
        font_paths = [
            "C:/Windows/Fonts/arial.ttf",
            "C:/Windows/Fonts/DejaVuSans.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/dejavu/DejaVuSans.ttf",
        ]
        for font_path in font_paths:
            if os.path.exists(font_path):
                pdfmetrics.registerFont(TTFont('CustomFont', font_path))
                font_name = 'CustomFont'
                break
    except Exception:
        pass

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=1.5*cm,
        leftMargin=1.5*cm,
        topMargin=1.5*cm,
        bottomMargin=1.5*cm
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='TitleRu', fontName=font_name, fontSize=16, leading=20, spaceAfter=12, textColor=colors.HexColor('#2c3e50')))
    styles.add(ParagraphStyle(name='HeadingRu', fontName=font_name, fontSize=12, leading=16, spaceBefore=10, spaceAfter=4, textColor=colors.HexColor('#34495e')))
    styles.add(ParagraphStyle(name='NormalRu', fontName=font_name, fontSize=9, leading=12, spaceAfter=4))
    styles.add(ParagraphStyle(name='TableCell', fontName=font_name, fontSize=8, leading=10))

    story = []
    story.append(Paragraph("Обезличенный документ", styles['TitleRu']))
    story.append(Paragraph(f"<b>Источник:</b> {source_name}", styles['NormalRu']))
    story.append(Spacer(1, 15))

    def parse_markdown_table(lines_block):
        table_data = []
        for line in lines_block:
            if '|' in line:
                if re.match(r'^[\s|:-]+$', line):
                    continue
                cells = [c.strip() for c in line.split('|')]
                if cells and cells[0] == '':
                    cells = cells[1:]
                if cells and cells[-1] == '':
                    cells = cells[:-1]
                if cells:
                    table_data.append(cells)
        return table_data

    def parse_html_table(html_text):
        table_data = []
        rows = re.findall(r'<tr[^>]*>(.*?)</tr>', html_text, re.DOTALL | re.IGNORECASE)
        for row in rows:
            cells = []
            for th in re.findall(r'<th[^>]*>(.*?)</th>', row, re.DOTALL | re.IGNORECASE):
                cell_text = re.sub(r'<[^>]+>', ' ', th).strip()
                cell_text = re.sub(r'\s+', ' ', cell_text)
                cells.append(cell_text[:80])
            for td in re.findall(r'<td[^>]*>(.*?)</td>', row, re.DOTALL | re.IGNORECASE):
                cell_text = re.sub(r'<[^>]+>', ' ', td).strip()
                cell_text = re.sub(r'\s+', ' ', cell_text)
                cells.append(cell_text[:80])
            if cells:
                table_data.append(cells)
        return table_data

    def add_table_to_story(table_data):
        if not table_data or len(table_data) < 1:
            return
        max_cols = max(len(row) for row in table_data)
        normalized_data = []
        for row in table_data:
            while len(row) < max_cols:
                row.append('')
            cell_row = [Paragraph(str(cell)[:100], styles['TableCell']) for cell in row[:max_cols]]
            normalized_data.append(cell_row)
        if not normalized_data:
            return
        available_width = A4[0] - 3*cm
        col_width = available_width / max_cols
        col_widths = [col_width] * max_cols
        try:
            table = Table(normalized_data, colWidths=col_widths)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498db')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('LEFTPADDING', (0, 0), (-1, -1), 4),
                ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#bdc3c7')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5f5f5')]),
            ]))
            story.append(table)
            story.append(Spacer(1, 10))
        except Exception:
            for row in table_data:
                story.append(Paragraph(' | '.join(str(c) for c in row), styles['NormalRu']))

    html_tables = re.findall(r'<table[^>]*>.*?</table>', html_content, re.DOTALL | re.IGNORECASE)
    table_placeholder_map = {}
    modified_content = html_content
    for idx, table_html in enumerate(html_tables):
        placeholder = f"__TABLE_PLACEHOLDER_{idx}__"
        table_placeholder_map[placeholder] = table_html
        modified_content = modified_content.replace(table_html, f"\n{placeholder}\n", 1)

    lines = modified_content.split('\n')
    current_table_lines = []

    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if line.startswith('__TABLE_PLACEHOLDER_'):
            placeholder = line
            if placeholder in table_placeholder_map:
                table_html = table_placeholder_map[placeholder]
                table_data = parse_html_table(table_html)
                if table_data:
                    add_table_to_story(table_data)
            i += 1
            continue

        if '|' in line and not line.startswith('<'):
            current_table_lines = [line]
            i += 1
            while i < len(lines):
                next_line = lines[i].strip()
                if '|' in next_line or re.match(r'^[\s:-]+$', next_line):
                    current_table_lines.append(next_line)
                    i += 1
                elif not next_line:
                    i += 1
                    continue
                else:
                    break
            table_data = parse_markdown_table(current_table_lines)
            if table_data:
                add_table_to_story(table_data)
            current_table_lines = []
            continue

        if not line:
            i += 1
            continue

        clean_line = re.sub(r'<[^>]+>', '', line)
        if not clean_line.strip():
            i += 1
            continue

        if clean_line.startswith('# ') or line.startswith('<h1'):
            story.append(Paragraph(clean_line.lstrip('# '), styles['TitleRu']))
        elif clean_line.startswith('## ') or line.startswith('<h2'):
            story.append(Paragraph(clean_line.lstrip('# '), styles['HeadingRu']))
        elif clean_line.startswith('### ') or line.startswith('<h3'):
            story.append(Paragraph(clean_line.lstrip('# '), styles['HeadingRu']))
        elif clean_line.startswith('---'):
            story.append(Spacer(1, 20))
        else:
            story.append(Paragraph(clean_line, styles['NormalRu']))

        i += 1

    if current_table_lines:
        table_data = parse_markdown_table(current_table_lines)
        if table_data:
            add_table_to_story(table_data)

    doc.build(story)


# ============ API Endpoints ============

@router.post("/upload")
async def upload_file(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    settings_str: str = Form(..., alias="settings")
):
    """Upload a file for anonymization."""
    ext = get_file_extension(file.filename)
    if ext not in ANONYMIZER_SUPPORTED_FORMATS:
        raise HTTPException(
            status_code=400,
            detail=f"Неподдерживаемый формат файла: {ext}. Поддерживаются: {', '.join(ANONYMIZER_SUPPORTED_FORMATS.keys())}"
        )

    content = await file.read()

    if len(content) > MAX_FILE_SIZE_BYTES:
        raise HTTPException(
            status_code=400,
            detail=f"Файл слишком большой. Максимальный размер: {MAX_FILE_SIZE_BYTES // (1024*1024)} МБ"
        )

    try:
        settings_dict = json.loads(settings_str)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Неверный формат настроек")

    task_id = str(uuid4())
    original_path = save_uploaded_file(content, file.filename, task_id)

    tasks[task_id] = {
        "status": "processing",
        "progress": 0,
        "message": "Начинаем обработку...",
        "filename": file.filename,
        "file_type": ext.replace(".", ""),
        "settings": settings_dict,
        "original_path": str(original_path),
        "created_at": datetime.now().isoformat(),
    }

    background_tasks.add_task(process_document, task_id)
    return {"task_id": task_id}


@router.get("/status/{task_id}")
async def get_status(task_id: str):
    """Get processing status for a task."""
    task = tasks.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Задача не найдена")

    return {
        "status": task["status"],
        "progress": task["progress"],
        "message": task["message"],
        "logs": task.get("logs", []),
    }


@router.get("/preview/{task_id}")
async def get_preview(task_id: str):
    """Get before/after preview for a task."""
    task = tasks.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Задача не найдена")

    if task["status"] != "done":
        raise HTTPException(status_code=400, detail="Обработка еще не завершена")

    return {
        "original": task.get("original_text", ""),
        "anonymized": task.get("anonymized_text", ""),
        "replacements_count": len(task.get("replacements", [])),
        "validation": task.get("validation", {}),
        "file_type": task.get("file_type", ""),
    }


@router.get("/download/{task_id}")
async def download_file(task_id: str):
    """Download anonymized file."""
    task = tasks.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Задача не найдена")

    if task["status"] != "done":
        raise HTTPException(status_code=400, detail="Обработка еще не завершена")

    output_path = Path(task["output_path"])
    if not output_path.exists():
        raise HTTPException(status_code=404, detail="Файл не найден")

    original_name = Path(task["filename"]).stem
    actual_ext = output_path.suffix
    download_name = f"{original_name}_anonymized{actual_ext}"

    media_types = {
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".pdf": "application/pdf",
        ".md": "text/markdown",
    }
    media_type = media_types.get(actual_ext, "application/octet-stream")

    return FileResponse(
        output_path,
        filename=download_name,
        media_type=media_type
    )


@router.get("/download-pdf/{task_id}")
async def download_pdf(task_id: str):
    """Download anonymized file as PDF."""
    task = tasks.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Задача не найдена")

    if task["status"] != "done":
        raise HTTPException(status_code=400, detail="Обработка еще не завершена")

    output_path = Path(task["output_path"])
    if not output_path.exists():
        raise HTTPException(status_code=404, detail="Файл не найден")

    content = ""
    if output_path.suffix == ".md":
        with open(output_path, 'r', encoding='utf-8') as f:
            content = f.read()
    else:
        content = task.get("anonymized_text", "")

    pdf_path = output_path.parent / f"{output_path.stem}.pdf"

    try:
        html_content = markdown_to_html(content)
        create_pdf_from_html(html_content, pdf_path, task["filename"])
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка создания PDF: {str(e)}")

    original_name = Path(task["filename"]).stem
    download_name = f"{original_name}_anonymized.pdf"

    return FileResponse(
        pdf_path,
        filename=download_name,
        media_type="application/pdf"
    )


@router.get("/view/{task_id}", response_class=HTMLResponse)
async def view_anonymized_file(task_id: str):
    """View anonymized file in a formatted HTML page."""
    task = tasks.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Задача не найдена")

    if task["status"] != "done":
        raise HTTPException(status_code=400, detail="Обработка еще не завершена")

    output_path = Path(task["output_path"])
    if not output_path.exists():
        raise HTTPException(status_code=404, detail="Файл не найден")

    content = ""
    if output_path.suffix == ".md":
        with open(output_path, 'r', encoding='utf-8') as f:
            content = f.read()
    elif output_path.suffix == ".docx":
        try:
            from docx import Document
            doc = Document(output_path)
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    paragraphs.append(" | ".join(cells))
            content = "\n\n".join(paragraphs)
        except Exception as e:
            content = task.get("anonymized_text", f"Ошибка чтения файла: {str(e)}")
    else:
        content = task.get("anonymized_text", "Предпросмотр недоступен для данного формата")

    html_content = markdown_to_html(content)

    # Use /api/v1/anonymizer/ prefix for links
    html_page = f"""<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <title>Просмотр: {task["filename"]}</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, sans-serif;
            background: #1a1a2e;
            color: #e0e0e0;
            line-height: 1.6;
            padding: 2rem;
        }}
        .container {{
            max-width: 900px;
            margin: 0 auto;
            background: #16213e;
            border-radius: 12px;
            padding: 2rem;
            box-shadow: 0 4px 20px rgba(0,0,0,0.3);
        }}
        h1 {{ color: #ffc107; margin-bottom: 1rem; font-size: 1.5rem; }}
        h2 {{ color: #4fc3f7; margin: 1.5rem 0 0.5rem; font-size: 1.2rem; }}
        h3 {{ color: #81c784; margin: 1rem 0 0.5rem; font-size: 1rem; }}
        p {{ margin-bottom: 1rem; }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 1rem 0;
            background: #1a1a3e;
            border-radius: 8px;
            overflow: hidden;
        }}
        th, td {{
            padding: 0.75rem;
            text-align: left;
            border-bottom: 1px solid #333;
        }}
        th {{
            background: #0f3460;
            color: #ffc107;
            font-weight: 600;
        }}
        tr:hover {{ background: #1e3a5f; }}
        ul, ol {{ margin-left: 1.5rem; margin-bottom: 1rem; }}
        li {{ margin-bottom: 0.5rem; }}
        hr {{ border: none; border-top: 1px solid #333; margin: 2rem 0; }}
        code {{ background: #0f3460; padding: 0.2rem 0.4rem; border-radius: 4px; }}
        .header {{
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 2rem;
            padding-bottom: 1rem;
            border-bottom: 1px solid #333;
        }}
        .badge {{
            background: #4caf50;
            color: white;
            padding: 0.25rem 0.75rem;
            border-radius: 999px;
            font-size: 0.875rem;
        }}
        .anonymized {{
            background: #ff9800;
            color: #000;
            padding: 0.1rem 0.3rem;
            border-radius: 3px;
            font-weight: 600;
        }}
        .toolbar {{
            position: fixed;
            top: 1rem;
            right: 1rem;
            display: flex;
            gap: 0.5rem;
            z-index: 100;
        }}
        .btn {{
            display: inline-flex;
            align-items: center;
            gap: 0.5rem;
            padding: 0.75rem 1.25rem;
            border: none;
            border-radius: 8px;
            cursor: pointer;
            font-size: 0.875rem;
            font-weight: 500;
            transition: all 0.2s;
            text-decoration: none;
        }}
        .btn-primary {{
            background: #f59e0b;
            color: #000;
        }}
        .btn-primary:hover {{
            background: #fbbf24;
        }}
        .btn-secondary {{
            background: #27272a;
            color: #fff;
            border: 1px solid #3f3f46;
        }}
        .btn-secondary:hover {{
            background: #3f3f46;
        }}
    </style>
</head>
<body>
    <div class="toolbar">
        <a href="/api/v1/anonymizer/download/{task_id}" class="btn btn-primary">⬇️ Скачать MD</a>
        <a href="/api/v1/anonymizer/download-pdf/{task_id}" class="btn btn-primary">📄 Скачать PDF</a>
        <button onclick="closeViewer()" class="btn btn-secondary">✕ Закрыть</button>
    </div>
    <script>
        function closeViewer() {{
            window.close();
            setTimeout(function() {{
                alert('Пожалуйста, закройте эту вкладку вручную, чтобы вернуться к результатам.');
            }}, 100);
        }}
    </script>
    <div class="container">
        <div class="header">
            <h1>📄 {task["filename"]}</h1>
            <span class="badge">🔒 Обезличено</span>
        </div>
        <div class="content">
            {html_content}
        </div>
    </div>
</body>
</html>"""

    return HTMLResponse(content=html_page)


@router.get("/mapping/{task_id}")
async def get_mapping(task_id: str):
    """Get anonymization mapping (replacements log)."""
    task = tasks.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Задача не найдена")

    if task["status"] != "done":
        raise HTTPException(status_code=400, detail="Обработка еще не завершена")

    mapping = {
        "file": task["filename"],
        "date": task["created_at"],
        "settings": task["settings"],
        "replacements": task.get("replacements", []),
        "validation": task.get("validation", {})
    }

    return JSONResponse(
        content=mapping,
        headers={
            "Content-Disposition": f"attachment; filename=mapping_{task_id[:8]}.json"
        }
    )


@router.post("/validate/{task_id}")
async def validate_result(task_id: str):
    """Run validation on anonymized result."""
    task = tasks.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Задача не найдена")

    if task["status"] != "done":
        raise HTTPException(status_code=400, detail="Обработка еще не завершена")

    anonymized_text = task.get("anonymized_text", "")

    try:
        validation_result = await validator.validate_with_ml(
            anonymized_text,
            task["settings"]
        )

        return {
            "is_valid": validation_result.is_valid,
            "issues": validation_result.issues,
            "warnings": validation_result.warnings,
            "confidence": validation_result.confidence
        }
    except Exception as e:
        return {
            "is_valid": task.get("validation", {}).get("is_valid", True),
            "issues": task.get("validation", {}).get("issues", []),
            "warnings": [f"ML валидация недоступна: {str(e)}"],
            "confidence": task.get("validation", {}).get("confidence", 0.5)
        }


@router.get("/profiles")
async def get_profiles():
    """Get available anonymization profiles."""
    return DEFAULT_PROFILES


@router.get("/ml-status")
async def get_ml_status():
    """Check ML models availability."""
    return ml_integration.is_available()


@router.delete("/task/{task_id}")
async def delete_task(task_id: str):
    """Delete a task and its files."""
    task = tasks.get(task_id)
    if task:
        delete_task_files(task_id)
        del tasks[task_id]

    return {"status": "deleted"}
