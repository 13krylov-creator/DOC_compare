"""FastAPI router for Document Analysis — 4th tab.

In-memory document analysis: upload, parse, count tokens/sheets,
search, Q&A, summarize and generate tables via LLM.
No persistent storage — documents live only in the tasks dict.
"""

import json
import re
import io
import asyncio
from pathlib import Path
from typing import Optional, List, Dict, Any
from datetime import datetime
from uuid import uuid4

import httpx
from fastapi import APIRouter, File, UploadFile, Form, HTTPException, Query, Response
from fastapi.responses import JSONResponse
from pydantic import BaseModel

try:
    import tiktoken
except ImportError:
    tiktoken = None

from config import ML_CONFIG

from anonymizer_core.ml_integration import MLIntegration

try:
    ml_integration = MLIntegration()
except Exception:
    ml_integration = None

router = APIRouter()

# --------------- in-memory task storage ---------------
tasks: Dict[str, Dict[str, Any]] = {}

# --------------- constants ---------------
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".txt", ".jpg", ".jpeg", ".png"}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB


# --------------- pydantic models ---------------
class AskRequest(BaseModel):
    question: str
    custom_prompt: str = ""


class SummarizeRequest(BaseModel):
    custom_prompt: str = ""


class TableRequest(BaseModel):
    custom_prompt: str = ""
    
class OCRRequest(BaseModel):
    pass


# --------------- helpers ---------------

def _count_tokens(text: str) -> int:
    """Accurate token count using tiktoken (cl100k_base)."""
    if not text:
        return 0
    if tiktoken:
        try:
            enc = tiktoken.get_encoding("cl100k_base")
            return len(enc.encode(text))
        except Exception:
            pass
    # Fallback if tiktoken fails or not installed
    return max(1, len(text) // 4)


def _call_gpt_sync(prompt: str, max_tokens: int = 4096, temperature: float = 0.2) -> str:
    """
    Synchronous GPT call via OpenAI-compatible API.
    Identical stack to MLIntegration.ask_gpt / LLMClient._llm_extract.
    """
    import requests as req
    import time

    host = ML_CONFIG["gpt"]["host"]
    model = ML_CONFIG["gpt"]["model"]
    last_error = ""

    for attempt in range(3):
        try:
            if attempt > 0:
                time.sleep(2 ** attempt)

            resp = req.post(
                f"http://{host}/v1/chat/completions",
                json={
                    "model": model,
                    "messages": [
                        {"role": "system", "content": "Ты — полезный и краткий помощник. Отвечай только на русском языке."},
                        {"role": "user", "content": prompt},
                    ],
                    "max_tokens": max_tokens,
                    "temperature": temperature,
                },
                timeout=180,
            )

            if resp.status_code == 200:
                data = resp.json()
                content = data["choices"][0]["message"]["content"]
                if content and len(content) > 2:
                    return content
                last_error = f"Пустой ответ GPT (попытка {attempt + 1})"
            elif resp.status_code == 429:
                last_error = f"GPT перегружен (429, попытка {attempt + 1})"
            elif resp.status_code >= 500:
                last_error = f"Ошибка сервера GPT ({resp.status_code})"
            else:
                last_error = f"HTTP {resp.status_code}: {resp.text[:120]}"
                break
        except req.exceptions.Timeout:
            last_error = f"Таймаут GPT (попытка {attempt + 1})"
        except req.exceptions.ConnectionError:
            last_error = f"GPT недоступен (попытка {attempt + 1})"
        except Exception as e:
            last_error = str(e)[:100]

    raise HTTPException(status_code=502, detail=f"LLM недоступна: {last_error}")


# --------------- document parsing ---------------

def _parse_pdf(file_bytes: bytes) -> dict:
    """Parse PDF: extract text per page."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise HTTPException(status_code=500, detail="PyMuPDF не установлен")

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    sheets = []
    for i, page in enumerate(doc):
        text = page.get_text()
        sheets.append({
            "name": f"Страница {i + 1}",
            "text": text,
            "tokens": _count_tokens(text),
        })
    doc.close()
    return {"sheets": sheets}


def _parse_docx(file_bytes: bytes) -> dict:
    """Parse DOCX: split by hard page breaks."""
    try:
        from docx import Document
        from docx.document import Document as _Document
        from docx.oxml.text.paragraph import CT_P
        from docx.oxml.table import CT_Tbl
        from docx.table import _Cell, Table
        from docx.text.paragraph import Paragraph
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx не установлен")

    doc = Document(io.BytesIO(file_bytes))
    
    # We will iterate over document body elements (paragraphs and tables) in order
    # and split content when we encounter a 'lastRenderedPageBreak' or soft break?
    # Actually python-docx doesn't easily show page layout. 
    # But hard breaks (`w:br` with type="page") are inside runs.
    
    # Strategy: Iterate paragraphs. Inside paragraphs, look for runs with breaks.
    
    pages = []
    current_page_text = []

    def flush_page():
        text = "\n".join(current_page_text).strip()
        if text:
            pages.append({
                "name": f"Страница {len(pages) + 1}",
                "text": text,
                "tokens": _count_tokens(text)
            })
        current_page_text.clear()

    # Access the underlying XML to iterate paragraphs and tables in order
    # This is tricky with python-docx. A simpler approach is to iterate paragraphs and tables
    # but we lose their relative order if we do them separately.
    # To keep order, we can iterate over `doc.element.body`.
    
    for element in doc.element.body:
        if isinstance(element, CT_P):
            paragraph = Paragraph(element, doc)
            para_text = ""
            
            # Check for page breaks in runs
            # We construct paragraph text, if we hit a break, we split
            
            # Simplified approach: if a paragraph contains a hard break, 
            # we consider the expected behavior (split before or after?)
            # Usually specific run has the break.
            
            parts = []
            has_break = False
            
            for run in paragraph.runs:
                # Check for explicit page break <w:br w:type="page"/>
                if 'w:br' in run._element.xml and 'w:type="page"' in run._element.xml:
                    # Flush current text so far as a page
                    if parts:
                        current_page_text.append("".join(parts))
                        parts = []
                    flush_page()
                    has_break = True
                
                parts.append(run.text)
            
            current_page_text.append("".join(parts))
            
            # Also check for 'pageBreakBefore' in paragraph format
            if paragraph.paragraph_format.page_break_before and not has_break:
                # If this paragraph starts a new page, flush PREVIOUS content first
                # But wait, we just appended this paragraph. 
                # Actually we should flush BEFORE appending this paragraph if it has page_break_before.
                # So let's backtrack slightly: pop the last added line? No, complicated.
                pass 
                # Improving logic: checks before appending
        
        elif isinstance(element, CT_Tbl):
            # It's a table
            table = Table(element, doc)
            # Flatten table to text
            rows_text = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_text.append(" | ".join(cells))
            current_page_text.append("\n" + "\n".join(rows_text) + "\n")

    flush_page()

    # If no explicit pages found (short doc or no breaks), treat as single page
    if not pages:
        full_text = "\n".join([p.text for p in doc.paragraphs])
        return {"sheets": [{"name": "Документ", "text": full_text, "tokens": _count_tokens(full_text)}]}

    return {"sheets": pages}


def _parse_xlsx(file_bytes: bytes) -> dict:
    """Parse XLSX/XLS: extract text per worksheet."""
    try:
        import openpyxl
    except ImportError:
        raise HTTPException(status_code=500, detail="openpyxl не установлен")

    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    sheets = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            rows.append(" | ".join(cells))
        text = "\n".join(rows)
        sheets.append({
            "name": sheet_name,
            "text": text,
            "tokens": _count_tokens(text),
        })
    wb.close()
    return {"sheets": sheets}


def _parse_txt(file_bytes: bytes) -> dict:
    """Parse TXT files."""
    for enc in ("utf-8", "cp1251", "latin-1"):
        try:
            text = file_bytes.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = file_bytes.decode("utf-8", errors="ignore")

    return {"sheets": [{"name": "Документ", "text": text, "tokens": _count_tokens(text)}]}


def _parse_document(file_bytes: bytes, ext: str) -> dict:
    """Route to the correct parser based on extension."""
    ext = ext.lower()
    if ext == ".pdf":
        return _parse_pdf(file_bytes)
    elif ext == ".docx" or ext == ".doc":
        return _parse_docx(file_bytes)
    elif ext in (".xlsx", ".xls"):
        return _parse_xlsx(file_bytes)
    elif ext == ".txt":
        return _parse_txt(file_bytes)
    elif ext in (".jpg", ".jpeg", ".png"):
        # For images, we just treat them as empty text/OCR candidates
        return {"sheets": [{"name": "Изображение", "text": "[Изображение для OCR]", "tokens": 0}]}
    else:
        raise HTTPException(status_code=400, detail=f"Неподдерживаемый формат: {ext}")


def _full_text(task: dict) -> str:
    """Get full concatenated text from all sheets."""
    return "\n\n".join(s["text"] for s in task["sheets"])


def _truncate_for_llm(text: str, max_chars: int = 24000) -> str:
    """Truncate text to fit within LLM context window."""
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + f"\n\n... (текст обрезан, показано {max_chars} из {len(text)} символов)"


# --------------- endpoints ---------------

@router.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    """Upload and parse a document. Returns task_id and stats."""
    # Validate extension
    filename = file.filename or "unknown"
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Неподдерживаемый формат: {ext}. Поддерживаются: {', '.join(ALLOWED_EXTENSIONS)}"
        )

    # Read file
    file_bytes = await file.read()
    if len(file_bytes) > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail="Файл слишком большой (макс. 50 МБ)")
    if len(file_bytes) == 0:
        raise HTTPException(status_code=400, detail="Файл пуст")

    # Parse
    parsed = _parse_document(file_bytes, ext)
    sheets = parsed["sheets"]
    total_tokens = sum(s["tokens"] for s in sheets)

    task_id = str(uuid4())
    tasks[task_id] = {
        "id": task_id,
        "filename": filename,
        "file_size": len(file_bytes),
        "ext": ext,
        "sheets": sheets,
        "total_tokens": total_tokens,
        "created_at": datetime.now().isoformat(),
        "file_bytes": file_bytes,  # Store for OCR
        "content_text": _full_text({"sheets": sheets}), # Store full text for analysis/editing
    }

    return {
        "task_id": task_id,
        "filename": filename,
        "file_size": len(file_bytes),
        "sheets_count": len(sheets),
        "sheets": [
            {"name": s["name"], "tokens": s["tokens"]}
            for s in sheets
        ],
        "total_tokens": total_tokens,
    }


@router.get("/info/{task_id}")
async def get_document_info(task_id: str):
    """Return cached document statistics."""
    task = tasks.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Документ не найден. Загрузите заново.")

    return {
        "task_id": task_id,
        "filename": task["filename"],
        "file_size": task["file_size"],
        "sheets_count": len(task["sheets"]),
        "sheets": [
            {"name": s["name"], "tokens": s["tokens"]}
            for s in task["sheets"]
        ],
        "total_tokens": task["total_tokens"],
    }


@router.post("/ask/{task_id}")
async def ask_document(task_id: str, body: AskRequest):
    """Ask a question about the document via LLM."""
    task = tasks.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Документ не найден. Загрузите заново.")

    question = body.question.strip()
    if not question:
        raise HTTPException(status_code=400, detail="Вопрос не указан")

    full = _full_text(task)
    doc_text = _truncate_for_llm(full)

    custom_part = ""
    if body.custom_prompt.strip():
        custom_part = f"\n\nДополнительные указания пользователя: {body.custom_prompt.strip()}"

    prompt = f"""Ты — эксперт-аналитик документов. Тебе предоставлен текст документа.
Ответь на вопрос пользователя на основе ТОЛЬКО этого документа. Если ответа нет в документе — честно скажи об этом.
Используй структурированный ответ: маркированные списки, жирный текст для ключевых моментов.{custom_part}

=== ДОКУМЕНТ ({task['filename']}) ===

{doc_text}

=== ВОПРОС ===

{question}"""

    answer = _call_gpt_sync(prompt, max_tokens=4096, temperature=0.2)

    return {
        "question": question,
        "answer": answer,
        "tokens_used": _count_tokens(doc_text),
    }


@router.post("/summarize/{task_id}")
async def summarize_document(task_id: str, body: SummarizeRequest):
    """Summarize / create a protocol of the document via LLM."""
    task = tasks.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Документ не найден. Загрузите заново.")

    full = _full_text(task)
    doc_text = _truncate_for_llm(full)

    custom_part = ""
    if body.custom_prompt.strip():
        custom_part = f"\n\nДополнительные указания пользователя по стилю или структуре суммаризации: {body.custom_prompt.strip()}"

    prompt = f"""Ты — профессиональный аналитик. Тебе предоставлен текст документа.

Создай подробный **протокол / суммаризацию** этого документа. Структура:

1. **Тип документа** — определи что это (договор, протокол, отчёт, письмо и т.д.)
2. **Краткое описание** — 2-3 предложения о чём документ
3. **Ключевые участники / стороны** — если есть
4. **Основные пункты / выводы** — маркированный список главных положений
5. **Даты и сроки** — если упомянуты
6. **Суммы и финансовые условия** — если упомянуты
7. **Итоги / решения** — если есть

Используй markdown для форматирования. Будь точен и конкретен.{custom_part}

=== ДОКУМЕНТ ({task['filename']}) ===

{doc_text}"""

    summary = _call_gpt_sync(prompt, max_tokens=4096, temperature=0.2)
    
    # Cache result for download
    task["summary"] = summary

    return {
        "summary": summary,
        "tokens_used": _count_tokens(doc_text),
    }


@router.get("/download/{task_id}/protocol")
async def download_protocol(task_id: str):
    """Generate DOCX from the cached summary."""
    task = tasks.get(task_id)
    if not task or "summary" not in task:
        raise HTTPException(status_code=404, detail="Протокол не найден. Сначала создайте его.")

    from docx import Document
    
    doc = Document()
    doc.add_heading(f"Протокол анализа: {task['filename']}", 0)
    
    # Simple markdown to docx conversion
    # Not full implementation, just paragraphs
    summary_text = task["summary"]
    
    for line in summary_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('# '):
            doc.add_heading(line.lstrip('# ').strip(), 1)
        elif line.startswith('## '):
            doc.add_heading(line.lstrip('# ').strip(), 2)
        elif line.startswith('### '):
            doc.add_heading(line.lstrip('# ').strip(), 3)
        elif line.startswith('- ') or line.startswith('* '):
             doc.add_paragraph(line.lstrip('-* ').strip(), style='List Bullet')
        elif re.match(r'^\d+\.', line):
             doc.add_paragraph(line, style='List Number')
        else:
            doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = f"Protocol_{task['filename']}.docx"
    
    # Handle UTF-8 filename in Content-Disposition
    from urllib.parse import quote
    encoded_filename = quote(filename)
    
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
        }
    )


@router.post("/table/{task_id}")
async def generate_table(task_id: str, body: TableRequest):
    """Generate structured tables."""
    task = tasks.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Документ не найден. Загрузите заново.")

    full = _full_text(task)
    doc_text = _truncate_for_llm(full)

    custom_part = ""
    if body.custom_prompt.strip():
        custom_part = f"""

КРИТИЧЕСКИ ВАЖНО — пользователь указал конкретные требования к таблице:
\"\"\"{body.custom_prompt.strip()}\"\"\"
Следуй этим указаниям максимально точно."""

    prompt = f"""Ты — эксперт по извлечению структурированных данных из документов.

Тебе предоставлен текст документа. Твоя задача — извлечь данные и представить их в виде ТАБЛИЦЫ.

ПРАВИЛА:
1. Проанализируй документ и определи, какие данные лучше всего представить таблицей
2. Выбери подходящие заголовки столбцов
3. Заполни строки данными из документа — ТОЛЬКО реальные данные, НЕ выдумывай
4. Если в документе несколько логических таблиц — создай несколько
5. Формат ответа — СТРОГО JSON, без markdown-обёртки

Ответь ТОЛЬКО валидным JSON в формате:
{{
  "tables": [
    {{
      "title": "Название таблицы",
      "headers": ["Столбец 1", "Столбец 2", "Столбец 3"],
      "rows": [
        ["Значение 1", "Значение 2", "Значение 3"],
        ["Значение 4", "Значение 5", "Значение 6"]
      ]
    }}
  ]
}}{custom_part}

=== ДОКУМЕНТ ({task['filename']}) ===

{doc_text}"""

    raw = _call_gpt_sync(prompt, max_tokens=4096, temperature=0.1)

    tables = []
    markdown_output = ""

    try:
        json_match = re.search(r'\{[\s\S]*\}', raw)
        if json_match:
            parsed = json.loads(json_match.group())
            raw_tables = parsed.get("tables", [])

            if not raw_tables and "headers" in parsed:
                raw_tables = [parsed]

            for tbl in raw_tables:
                title = tbl.get("title", "Таблица")
                headers = tbl.get("headers", [])
                rows = tbl.get("rows", [])

                if not headers and rows:
                    headers = [f"Столбец {i+1}" for i in range(len(rows[0]))]

                if not headers:
                    continue

                normalized_rows = []
                for row in rows:
                    if isinstance(row, list):
                        norm = [str(cell) if cell is not None else "" for cell in row]
                        while len(norm) < len(headers):
                            norm.append("")
                        norm = norm[:len(headers)]
                        normalized_rows.append(norm)
                    elif isinstance(row, dict):
                        norm = [str(row.get(h, "")) for h in headers]
                        normalized_rows.append(norm)

                tables.append({
                    "title": title,
                    "headers": headers,
                    "rows": normalized_rows,
                })

                md = f"### {title}\n\n"
                md += "| " + " | ".join(headers) + " |\n"
                md += "|" + "|".join(["---" for _ in headers]) + "|\n"
                for row in normalized_rows:
                    md += "| " + " | ".join(row) + " |\n"
                markdown_output += md + "\n"
        else:
            markdown_output = raw
            tables = []

    except Exception:
        markdown_output = raw
        tables = []

    # Cache tables for download
    task["tables"] = tables

    return {
        "tables": tables,
        "markdown": markdown_output,
        "raw_response": raw,
        "tokens_used": _count_tokens(doc_text),
    }


@router.get("/download/{task_id}/table")
async def download_table(task_id: str):
    """Generate XLSX from the cached tables."""
    task = tasks.get(task_id)
    if not task or "tables" not in task or not task["tables"]:
        raise HTTPException(status_code=404, detail="Таблицы не найдены. Сначала сгенерируйте их.")

    try:
        import openpyxl
        from openpyxl.utils import get_column_letter
    except ImportError:
        raise HTTPException(status_code=500, detail="openpyxl не установлен")

    wb = openpyxl.Workbook()
    # Remove default sheet
    default_ws = wb.active
    wb.remove(default_ws)

    for i, tbl in enumerate(task["tables"]):
        title = tbl.get("title", f"Table {i+1}")
        # Sanitize sheet title (max 31 chars, no specials)
        safe_title = re.sub(r'[\\/*?:\[\]]', '', title)[:30] or f"Sheet{i+1}"
        
        ws = wb.create_sheet(title=safe_title)
        
        # Headers
        ws.append(tbl.get("headers", []))
        
        # Rows
        for row in tbl.get("rows", []):
            ws.append(row)

        # Basic auto-width (approximate)
        for col_idx, col_cells in enumerate(ws.columns, 1):
            max_len = 0
            for cell in col_cells:
                try:
                    if len(str(cell.value)) > max_len:
                        max_len = len(str(cell.value))
                except: pass
            
            # Cap width
            adjusted_width = min(max_len + 2, 50)
            ws.column_dimensions[get_column_letter(col_idx)].width = adjusted_width

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    filename = f"Tables_{task['filename']}.xlsx"
    from urllib.parse import quote
    encoded_filename = quote(filename)

    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
        }
    )



# --------------- exporters ---------------

import re
from pathlib import Path

def _markdown_to_html(md_text: str) -> str:
    """Convert markdown to HTML with basic formatting."""
    html = md_text

    html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
    html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
    html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)

    html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)
    html = re.sub(r'__(.+?)__', r'<strong>\1</strong>', html)
    html = re.sub(r'\*(.+?)\*', r'<em>\1</em>', html)

    html = re.sub(r'^---+$', r'<hr>', html, flags=re.MULTILINE)
    html = re.sub(r'^- (.+)$', r'<li>\1</li>', html, flags=re.MULTILINE)
    html = re.sub(r'^(\d+)\. (.+)$', r'<li>\2</li>', html, flags=re.MULTILINE)

    lines = html.split('\n')
    result_lines = []
    for line in lines:
        line = line.strip()
        if line and not line.startswith('<') and not line.startswith('|'):
            line = f'<p>{line}</p>'
        result_lines.append(line)
    html = '\n'.join(result_lines)

    html = _convert_md_tables(html)
    return html


def _convert_md_tables(text: str) -> str:
    """Convert markdown tables to HTML."""
    lines = text.split('\n')
    result_lines = []
    in_table = False
    table_rows = []

    for line in lines:
        if '|' in line and not line.strip().startswith('<'):
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if cells:
                if all(c.replace('-', '').replace(':', '') == '' for c in cells):
                    continue
                if not in_table:
                    in_table = True
                    table_rows = []
                table_rows.append(cells)
        else:
            if in_table:
                result_lines.append(_render_table(table_rows))
                in_table = False
                table_rows = []
            result_lines.append(line)

    if in_table:
        result_lines.append(_render_table(table_rows))

    return '\n'.join(result_lines)


def _render_table(rows: list) -> str:
    """Render table rows as HTML."""
    if not rows:
        return ''

    html = '<table>\n'
    if rows:
        html += '<thead><tr>'
        for cell in rows[0]:
            html += f'<th>{cell}</th>'
        html += '</tr></thead>\n'

    if len(rows) > 1:
        html += '<tbody>\n'
        for row in rows[1:]:
            html += '<tr>'
            for cell in row:
                html += f'<td>{cell}</td>'
            html += '</tr>\n'
        html += '</tbody>\n'

    html += '</table>'
    return html


def _set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "000000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def _markdown_to_docx(md_text: str, output_stream):
    """HTML-based Markdown to DOCX converter with table support."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Convert to HTML first using the robust parser
    html_content = _markdown_to_html(md_text)
    
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Extract tables
    html_tables = re.findall(r'<table[^>]*>.*?</table>', html_content, re.DOTALL | re.IGNORECASE)
    table_placeholder_map = {}
    modified_content = html_content
    for idx, table_html in enumerate(html_tables):
        placeholder = f"__TABLE_PLACEHOLDER_{idx}__"
        table_placeholder_map[placeholder] = table_html
        modified_content = modified_content.replace(table_html, f"\n{placeholder}\n", 1)

    lines = modified_content.split('\n')
    
    # Helper for borders
    border_style = {"sz": 4, "val": "single", "color": "000000", "space": "0"}

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        if line.startswith('__TABLE_PLACEHOLDER_'):
            placeholder = line
            if placeholder in table_placeholder_map:
                table_html = table_placeholder_map[placeholder]
                # Parse HTML table
                rows = re.findall(r'<tr[^>]*>(.*?)</tr>', table_html, re.DOTALL | re.IGNORECASE)
                rows_data = []
                for row in rows:
                    cells = []
                    # Try to get th first, then td
                    th_cells = re.findall(r'<th[^>]*>(.*?)</th>', row, re.DOTALL | re.IGNORECASE)
                    td_cells = re.findall(r'<td[^>]*>(.*?)</td>', row, re.DOTALL | re.IGNORECASE)
                    
                    # Combine logical cells
                    raw_cells = th_cells + td_cells
                    
                    for cell_html in raw_cells:
                        # Strip HTML tags from cell content
                        cell_text = re.sub(r'<[^>]+>', ' ', cell_html).strip()
                        cell_text = re.sub(r'\s+', ' ', cell_text) # Normalize whitespace
                        cells.append(cell_text)
                        
                    if cells:
                        rows_data.append(cells)
                
                if rows_data:
                     max_cols = max(len(r) for r in rows_data)
                     if max_cols > 0:
                        table = doc.add_table(rows=len(rows_data), cols=max_cols)
                        table.style = 'Table Grid'
                        
                        for r_idx, row_data in enumerate(rows_data):
                            row = table.rows[r_idx]
                            for c_idx, cell_data in enumerate(row_data):
                                if c_idx < len(row.cells):
                                    cell = row.cells[c_idx]
                                    cell.text = cell_data
                                    
                                    # Apply borders to every cell
                                    _set_cell_border(
                                        cell, 
                                        top=border_style, bottom=border_style, 
                                        start=border_style, end=border_style,
                                        insideV=border_style, insideH=border_style
                                    )
                                    
                                    # Bold headers (first row)
                                    if r_idx == 0:
                                        for paragraph in cell.paragraphs:
                                            for run in paragraph.runs:
                                                run.bold = True
                        
                        doc.add_paragraph() # Spacer
            i += 1
            continue

        # Strip HTML tags for normal text
        clean_line = re.sub(r'<[^>]+>', '', line).strip()
        if not clean_line:
            i += 1
            continue
            
        if '<h1>' in line:
            # doc.add_heading(clean_line, 1) - cleaner to use styles
             p = doc.add_paragraph(clean_line)
             p.style = 'Heading 1'
        elif '<h2>' in line:
             p = doc.add_paragraph(clean_line)
             p.style = 'Heading 2'
        elif '<li>' in line:
             doc.add_paragraph(clean_line, style='List Bullet')
        else:
             doc.add_paragraph(clean_line)
        
        i += 1

    doc.save(output_stream)


def _create_pdf_from_html(html_content: str, output_stream):
    """Create PDF from markdown/HTML content using reportlab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import os

    font_name = 'Helvetica'
    try:
        font_paths = [
            "C:/Windows/Fonts/arial.ttf",
            "C:/Windows/Fonts/calibri.ttf", 
            "C:/Windows/Fonts/tahoma.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/dejavu/DejaVuSans.ttf",
        ]
        for font_path in font_paths:
            if os.path.exists(font_path):
                pdfmetrics.registerFont(TTFont('CustomFont', font_path))
                font_name = 'CustomFont'
                break
    except Exception:
        pass

    doc = SimpleDocTemplate(
        output_stream,
        pagesize=A4,
        rightMargin=1.5*cm,
        leftMargin=1.5*cm,
        topMargin=1.5*cm,
        bottomMargin=1.5*cm
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='TitleRu', fontName=font_name, fontSize=16, leading=20, spaceAfter=12, textColor=colors.HexColor('#2c3e50')))
    styles.add(ParagraphStyle(name='HeadingRu', fontName=font_name, fontSize=12, leading=16, spaceBefore=10, spaceAfter=4, textColor=colors.HexColor('#34495e')))
    styles.add(ParagraphStyle(name='NormalRu', fontName=font_name, fontSize=9, leading=12, spaceAfter=4))
    styles.add(ParagraphStyle(name='TableCell', fontName=font_name, fontSize=8, leading=10))

    story = []
    
    # Custom HTML parser logic (copied and adapted from Anonymizer)
    def parse_html_table(html_text):
        table_data = []
        rows = re.findall(r'<tr[^>]*>(.*?)</tr>', html_text, re.DOTALL | re.IGNORECASE)
        for row in rows:
            cells = []
            # th
            for th in re.findall(r'<th[^>]*>(.*?)</th>', row, re.DOTALL | re.IGNORECASE):
                cell_text = re.sub(r'<[^>]+>', ' ', th).strip()
                cells.append(cell_text[:80])
            # td
            for td in re.findall(r'<td[^>]*>(.*?)</td>', row, re.DOTALL | re.IGNORECASE):
                cell_text = re.sub(r'<[^>]+>', ' ', td).strip()
                cells.append(cell_text[:80])
            if cells:
                table_data.append(cells)
        return table_data

    def add_table_to_story(table_data):
        if not table_data or len(table_data) < 1:
            return
        
        # Normalize columns
        max_cols = max(len(row) for row in table_data)
        normalized_data = []
        for row in table_data:
            while len(row) < max_cols:
                row.append('')
            cell_row = [Paragraph(str(cell)[:100], styles['TableCell']) for cell in row[:max_cols]]
            normalized_data.append(cell_row)
            
        if not normalized_data:
            return
            
        available_width = A4[0] - 3*cm
        col_width = available_width / max_cols
        col_widths = [col_width] * max_cols
        try:
            table = Table(normalized_data, colWidths=col_widths)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498db')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('LEFTPADDING', (0, 0), (-1, -1), 4),
                ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#bdc3c7')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5f5f5')]),
            ]))
            story.append(table)
            story.append(Spacer(1, 10))
        except Exception:
             pass

    # Extract tables first
    html_tables = re.findall(r'<table[^>]*>.*?</table>', html_content, re.DOTALL | re.IGNORECASE)
    table_placeholder_map = {}
    modified_content = html_content
    for idx, table_html in enumerate(html_tables):
        placeholder = f"__TABLE_PLACEHOLDER_{idx}__"
        table_placeholder_map[placeholder] = table_html
        modified_content = modified_content.replace(table_html, f"\n{placeholder}\n", 1)

    lines = modified_content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        if line.startswith('__TABLE_PLACEHOLDER_'):
            placeholder = line
            if placeholder in table_placeholder_map:
                table_html = table_placeholder_map[placeholder]
                table_data = parse_html_table(table_html)
                if table_data:
                    add_table_to_story(table_data)
            i += 1
            continue

        clean_line = re.sub(r'<[^>]+>', '', line).strip()
        if not clean_line:
            i += 1
            continue
            
        if '<h1>' in line or '<h2>' in line:
             story.append(Paragraph(clean_line, styles['HeadingRu']))
        else:
             story.append(Paragraph(clean_line, styles['NormalRu']))
        
        i += 1
        
    doc.build(story)


def _markdown_to_pdf(md_text: str, output_stream):
    """Wrapper: Markdown -> HTML -> PDF."""
    html = _markdown_to_html(md_text)
    _create_pdf_from_html(html, output_stream)



@router.post("/ocr/{task_id}")
async def ocr_document(task_id: str):
    """OCR document to Markdown with tables."""
    task = tasks.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Документ не найден")
    
    if "file_bytes" not in task:
        raise HTTPException(status_code=400, detail="Файл удален или не доступен")
        
    if not ml_integration:
        raise HTTPException(status_code=503, detail="Система OCR не настроена")

    file_bytes = task["file_bytes"]
    filename = task["filename"]
    ext = task["ext"].lower()
    
    markdown_result = ""
    loop = asyncio.get_event_loop()
    
    import time
    start_time = time.time()
    pages_count = 0

    try:
        if ext == ".pdf":
            # Convert PDF pages to images
            import fitz
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            parts = []
            
            pages_count = len(doc)
            
            for i, page in enumerate(doc):
                # High resolution for OCR
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img_bytes = pix.tobytes("png")
                
                # Call Chandra sync in thread
                page_md = await loop.run_in_executor(None, ml_integration.ocr_with_chandra, img_bytes)
                if page_md:
                    parts.append(f"## Страница {i+1}\n\n{page_md}")
            
            doc.close()
            markdown_result = "\n\n---\n\n".join(parts)
            
        elif ext in [".jpg", ".jpeg", ".png"]:
            pages_count = 1
            markdown_result = await loop.run_in_executor(None, ml_integration.ocr_with_chandra, file_bytes)
        else:
            raise HTTPException(status_code=400, detail="OCR поддерживается только для PDF и изображений")
            
        if not markdown_result:
            raise HTTPException(status_code=500, detail="Не удалось распознать текст")
            
        task["ocr_markdown"] = markdown_result
        
        processing_time = time.time() - start_time
        
        return {
            "status": "success",
            "message": "Распознавание завершено",
            "preview": markdown_result[:500] + "...",
            "tokens_used": _count_tokens(markdown_result),
            "pages_count": pages_count,
            "processing_time": round(processing_time, 2)
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка OCR: {str(e)}")


@router.get("/download/{task_id}/ocr_docx")
async def download_ocr_docx(task_id: str):
    """Download OCR result as DOCX."""
    task = tasks.get(task_id)
    if not task or "ocr_markdown" not in task:
        raise HTTPException(status_code=404, detail="Результат OCR не найден")

    buffer = io.BytesIO()
    try:
        _markdown_to_docx(task["ocr_markdown"], buffer)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка генерации DOCX: {str(e)}")
        
    buffer.seek(0)
    
    filename = f"OCR_{task['filename']}.docx"
    from urllib.parse import quote
    encoded_filename = quote(filename)

    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
        }
    )


@router.get("/download/{task_id}/ocr_pdf")
async def download_ocr_pdf(task_id: str):
    """Download OCR result as PDF."""
    task = tasks.get(task_id)
    if not task or "ocr_markdown" not in task:
        raise HTTPException(status_code=404, detail="Результат OCR не найден")

    buffer = io.BytesIO()
    try:
        _markdown_to_pdf(task["ocr_markdown"], buffer)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка генерации PDF: {str(e)}")
        
    buffer.seek(0)
    
    filename = f"OCR_{task['filename']}.pdf"
    from urllib.parse import quote
    encoded_filename = quote(filename)

    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
        }
    )

class EditRequest(BaseModel):
    mode: str = "fix" # fix, style, toc, paraphrase
    custom_prompt: Optional[str] = None

@router.post("/edit/{task_id}")
async def edit_document(task_id: str, request: EditRequest):
    """Edit document text (Fix errors, Style, TOC, Paraphrase)."""
    if task_id not in tasks:
        raise HTTPException(status_code=404, detail="Task not found")
    
    task = tasks[task_id]
    
    # Use OCR markdown if available, otherwise original content
    content = task.get("ocr_markdown") or task.get("content_text") or task.get("markdown")
    
    if not content:
        # Fallback: try to reconstruct from sheets
        content = _full_text(task)
        
    if not content:
        raise HTTPException(status_code=400, detail="Не найден текст для редактирования. Возможно, это скан или изображение - попробуйте сначала запустить распознавание (OCR).")

    import time
    start_time = time.time()

    prompt = ""
    # system_role is not used directly in _call_gpt_sync (it uses hardcoded system prompt), 
    # so we should include it in the prompt or modify _call_gpt_sync.
    # _call_gpt_sync has hardcoded system: "Ты — полезный и краткий помощник..."
    # We need to override it? No, existing _call_gpt_sync doesn't support custom system role easily without modifying it.
    # However, we can just use the prompt.
    
    base_prompt = "Ты профессиональный редактор документов. Твоя задача - обработать текст, сохранив его структуру, заголовки и таблицы (в Markdown).\n\n"

    if request.mode == "fix":
        prompt = base_prompt + "Исправь все орфографические, пунктуационные и грамматические ошибки. Не меняй смысл и стиль, только исправь ошибки."
    elif request.mode == "style":
        prompt = base_prompt + "Улучши стиль текста, сделав его более официальным, профессиональным и читаемым. Сохрани структуру."
    elif request.mode == "paraphrase":
        prompt = base_prompt + "Перефразируй текст для повышения уникальности, сохраняя исходный смысл и все факты. Используй синонимы и меняй структуру предложений."
    elif request.mode == "toc":
        prompt = base_prompt + "Сгенерируй оглавление для этого текста и вставь его в начало. Используй ссылки-якоря если возможно, или просто список. Сам текст менять не нужно, просто добавь оглавление в начало."
    else:
        prompt = base_prompt + "Обработай документ согласно инструкции."

    if request.custom_prompt:
        prompt += f"\nДополнительные указания: {request.custom_prompt}"

    prompt += f"\n\n=== ТЕКСТ ДЛЯ ОБРАБОТКИ ===\n\n{content}\n\n=== КОНЕЦ ТЕКСТА ===\n\nВАЖНО: Верни ПОЛНЫЙ текст документа в формате Markdown. Обязательно сохрани все таблицы!"

    try:
        # We process the whole text (might be large, using large context model)
        # Using _call_gpt_sync which is available in this file
        updated_text = _call_gpt_sync(prompt, max_tokens=4096, temperature=0.2)
        
        task["edit_markdown"] = updated_text
        task["edit_mode"] = request.mode
        
        processing_time = time.time() - start_time
        
        # Generate diff
        import difflib
        
        def _generate_inline_diff(old, new):
            """Generate inline HTML diff."""
            # Helper to tokenize by words but keep whitespace
            def tokenize(text):
                return re.findall(r'\S+|\s+', text)

            a = tokenize(old)
            b = tokenize(new)
            
            matcher = difflib.SequenceMatcher(None, a, b)
            html = []
            
            for opcode, a0, a1, b0, b1 in matcher.get_opcodes():
                if opcode == 'equal':
                    html.append("".join(a[a0:a1]))
                elif opcode == 'insert':
                    html.append(f'<span style="background-color: #d4edda; color: #155724; text-decoration: none;">{"".join(b[b0:b1])}</span>')
                elif opcode == 'delete':
                    html.append(f'<span style="background-color: #f8d7da; color: #721c24; text-decoration: line-through; opacity: 0.7;">{"".join(a[a0:a1])}</span>')
                elif opcode == 'replace':
                    # show del then ins
                    html.append(f'<span style="background-color: #f8d7da; color: #721c24; text-decoration: line-through; opacity: 0.7;">{"".join(a[a0:a1])}</span>')
                    html.append(f'<span style="background-color: #d4edda; color: #155724; text-decoration: none;">{"".join(b[b0:b1])}</span>')
            
            return "".join(html)

        diff_html = _generate_inline_diff(content, updated_text)
        
        return {
            "status": "success",
            "preview": updated_text, # Return FULL text now (or at least larger chunk)
            "diff_view": diff_html, # New field for frontend
            "tokens_used": _count_tokens(updated_text),
            "processing_time": round(processing_time, 2)
        }
    except Exception as e:
        # logger.error(f"Edit error: {e}") # Logger not defined
        print(f"Edit error: {e}")
        raise HTTPException(status_code=500, detail=f"Ошибка обработки: {str(e)}")


@router.get("/download/{task_id}/edit_docx")
async def download_edit_docx(task_id: str):
    """Download edited document as DOCX."""
    task = tasks.get(task_id)
    if not task or "edit_markdown" not in task:
        raise HTTPException(status_code=404, detail="Result not found")
    
    md_text = task["edit_markdown"]
    filename = f"Edited_{task['filename']}.docx"
    from urllib.parse import quote
    
    from io import BytesIO
    output = BytesIO()
    _markdown_to_docx(md_text, output)
    output.seek(0)
    
    return Response(
        content=output.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"}
    )

@router.get("/download/{task_id}/edit_pdf")
async def download_edit_pdf(task_id: str):
    """Download edited document as PDF."""
    task = tasks.get(task_id)
    if not task or "edit_markdown" not in task:
        raise HTTPException(status_code=404, detail="Result not found")
    
    md_text = task["edit_markdown"]
    filename = f"Edited_{task['filename']}.pdf"
    from urllib.parse import quote
    
    from io import BytesIO
    output = BytesIO()
    _markdown_to_pdf(md_text, output)
    output.seek(0)
    
    return Response(
        content=output.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"}
    )


class StructureRequest(BaseModel):
    mode: str = "mindmap" # mindmap, flowchart, graph
    custom_prompt: Optional[str] = None

@router.post("/structure/{task_id}")
async def generate_structure_code(task_id: str, request: StructureRequest):
    """Generate Mermaid.js code for diagram."""
    if task_id not in tasks:
        raise HTTPException(status_code=404, detail="Task not found")
    
    task = tasks[task_id]
     # Use OCR markdown if available, otherwise original content
    content = task.get("ocr_markdown") or task.get("content_text") or task.get("markdown")
    
    if not content:
        content = _full_text(task)
        
    if not content:
        raise HTTPException(status_code=400, detail="Не найден текст для анализа.")

    import time
    start_time = time.time()

    prompt = ""
    system_role = "Ты эксперт по визуализации данных. Твоя задача - создать код диаграммы Mermaid.js на основе текста."

    base_prompt = "Проанализируй текст и создай структуру на языке Mermaid.js. Верни ТОЛЬКО код. НЕ используй markdown обертку (```). \nВАЖНО: \n1. Избегай спецсимволов в тексте узлов (скобки, кавычки) или экранируй их.\n2. Текст в узлах должен быть кратким.\n3. Если используешь graph/flowchart, текст в узлах бери в кавычки, например: id[\"Текст\"].\n"
    
    if request.mode == "mindmap":
        prompt = base_prompt + "Создай ментальную карту (mindmap). Синтаксис:\nmindmap\n  root((Тема))\n    Ветвь 1\n      Подветвь\n"
    elif request.mode == "flowchart":
        prompt = base_prompt + "Создай блок-схему (flowchart TD). Используй id[\"Текст\"] для узлов. Связи -->."
    elif request.mode == "graph":
        prompt = base_prompt + "Создай граф (graph TD). Выдели сущности и связи."
    else:
        prompt = base_prompt + "Создай диаграмму."

    if request.custom_prompt:
        prompt += f"\nДополнительные указания: {request.custom_prompt}"

    prompt += f"\n\n=== ТЕКСТ ===\n\n{content[:20000]}\n\n=== КОНЕЦ ТЕКСТА ===\n\nВерни ТОЛЬКО валидный код Mermaid."

    try:
        mermaid_code = _call_gpt_sync(prompt, max_tokens=2048, temperature=0.2)
        
        # Clean up code if LLM adds markdown wrapper
        mermaid_code = mermaid_code.replace("```mermaid", "").replace("```", "").strip()
        
        task["mermaid_code"] = mermaid_code
        task["structure_mode"] = request.mode
        
        processing_time = time.time() - start_time
        
        return {
            "status": "success",
            "mermaid_code": mermaid_code,
            "tokens_used": _count_tokens(content),
            "processing_time": round(processing_time, 2)
        }
    except Exception as e:
        print(f"Structure error: {e}")
        raise HTTPException(status_code=500, detail=f"Ошибка генерации: {str(e)}")

@router.delete("/{task_id}")
async def delete_task(task_id: str):
    """Remove document from memory."""
    if task_id in tasks:
        del tasks[task_id]
        return {"status": "deleted"}
    raise HTTPException(status_code=404, detail="Документ не найден")
